import tkinter as tk
from tkinter import ttk, messagebox
import tkinter.font as tkfont
from openpyxl import load_workbook
from openpyxl.styles import Font
from docxtpl import DocxTemplate
import docx
from datetime import datetime
import os
import glob
import json
import sqlite3

# ─────────────────────────────────────────────
#  CONFIGURATION — edit these paths as needed
# ─────────────────────────────────────────────
BASE_DIR        = r'C:\Users\RaDAR and LiDAR Lab\Documents\certificate'
HISTORY_DIR     = os.path.join(BASE_DIR, 'history')
AREA_CODE_FILE  = os.path.join(BASE_DIR, 'area_code.docx')
EXCEL_2026      = os.path.join(BASE_DIR, 'week_report_2026.xlsx')
TEMPLATE_DIR    = os.path.join(BASE_DIR, 'template')
RADAR_OUT_DIR   = os.path.join(BASE_DIR, 'radar')
LIDAR_OUT_DIR   = os.path.join(BASE_DIR, 'lidar')

SESSION_FILE    = os.path.join(BASE_DIR, 'session_units.json')
HISTORY_DB_FILE = os.path.join(BASE_DIR, 'history_index.sqlite3')
SEAL_IMAGE      = os.path.join(BASE_DIR, 'seal.png')
SEALED_DIR      = os.path.join(BASE_DIR, 'sealed')

RADAR_PREFIX    = 'AS26-'
LIDAR_PREFIX    = 'ASL26-'

RADAR_TYPES = ["DS", "DH", "DSP", "DE", "DEP", "AS", "ASP", "ZC", "ZM", "DC"]
LIDAR_TYPES = ["TS", "TJ_SXB", "TJ_S", "UX", "LP", "UL"]

# Persist last selected unit type across form clears
_last_radar_type = RADAR_TYPES[0]
_last_lidar_type = LIDAR_TYPES[0]

RADAR_TEMPLATES = {
    "DS":  "template_ds.docx",  "DH":  "template_dh.docx",
    "DSP": "template_ds_passed.docx", "DE":  "template_de.docx",
    "DEP": "template_de_passed.docx", "AS":  "template_as.docx",
    "ASP": "template_as_passed.docx", "ZC":  "template_zc.docx",
    "ZM":  "template_zm.docx",  "DC":  "template_dc.docx",
}
RADAR_UNIT_NAMES = {
    "DS": "Stalker DSR", "DH": "Stalker DSR", "DSP": "Stalker DSR",
    "DE": "Stalker DSR", "DEP": "Stalker DSR", "AS": "Stalker II SDR",
    "ASP": "Stalker II SDR", "ZC": "Stalker dual SL",
    "ZM": "Stalker dual SL", "DC": "Stalker dual SL",
}
LIDAR_TEMPLATES = {
    "TS":     "template_ts.docx",
    "TJ_SXB": "template_tj_sxb.docx",
    "TJ_S":   "template_tj_s.docx",
    "UX":     "template_ux.docx",
    "LP":     "template_lp.docx",
    "UL":     "template_ul.docx",
}
LIDAR_UNIT_NAMES = {
    "TS":     "LTI 20/20 TruSpeed LR",
    "TJ_SXB": "LTI 20/20 TruSpeed Sxb",
    "TJ_S":   "LTI 20/20 TruSpeed S",
    "UX":     "20/20 Ultralyte 200 LR",
    "LP":     "PRO-LITE+",
    "UL":     "20/20 Ultralyte 200 LR",
}

DEFAULT_FA = "326655"
DEFAULT_FB = "438384"

RETEST_PROMPT_YEARS = 2.5
RETEST_MIN_YEARS = 3.0

# ─────────────────────────────────────────────
#  COLORS & STYLES
# ─────────────────────────────────────────────
BG          = "#FFFFFF"
PANEL_BG    = "#F5F5F5"
CARD_BG     = "#F0E6E8"
ACCENT      = "#D41736"
GREEN       = "#00A39D"
AMBER       = "#B8860B"
TEXT        = "#1a1a1a"
TEXT_DIM    = "#6B6B6B"
TEXT_DARK   = "#FFFFFF"
WHITE       = "#FFFFFF"
RED         = "#A6192E"
BORDER      = "#D9D9D9"
ENTRY_BG    = "#FFFFFF"
LOCKED_BG   = "#E6F7F6"
LOCKED_FG   = "#007A76"


# ─────────────────────────────────────────────
#  HELPER FUNCTIONS
# ─────────────────────────────────────────────
def load_area_codes():
    content = docx.Document(AREA_CODE_FILE)
    return [p.text for p in content.paragraphs]

def get_last_lab_number(worksheet, prefix):
    """Return the highest lab number found in column A for the given prefix."""
    highest = 0
    for row in range(1, worksheet.max_row + 1):
        cell_value = worksheet.cell(row=row, column=1).value
        if cell_value and prefix in str(cell_value):
            try:
                clean = str(cell_value).replace(' (RETEST)', '').strip()
                num = int(clean.split('-')[1])
                if num > highest:
                    highest = num
            except (IndexError, ValueError):
                continue
    return highest

# ─────────────────────────────────────────────
#  HISTORY INDEX
# ─────────────────────────────────────────────
_HISTORY_INDEX = {'RADAR': {}, 'LIDAR': {}}
_INDEX_READY   = False
_INDEX_STATUS  = "Not built"

def _history_workbooks():
    pattern = os.path.join(HISTORY_DIR, 'week_report_*.xlsx')
    history_files = sorted(glob.glob(pattern), reverse=True)
    return history_files + [EXCEL_2026]

def _file_stamp(filepath):
    st = os.stat(filepath)
    return st.st_mtime_ns, st.st_size

def _history_db_connect():
    os.makedirs(BASE_DIR, exist_ok=True)
    conn = sqlite3.connect(HISTORY_DB_FILE, timeout=30)
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("PRAGMA synchronous=NORMAL")
    conn.execute("""
        CREATE TABLE IF NOT EXISTS files (
            path TEXT PRIMARY KEY,
            mtime_ns INTEGER NOT NULL,
            size INTEGER NOT NULL,
            file_order INTEGER NOT NULL,
            indexed_at TEXT NOT NULL
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS records (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            source_path TEXT NOT NULL,
            source_file TEXT NOT NULL,
            file_order INTEGER NOT NULL,
            row_order INTEGER NOT NULL,
            sheet_name TEXT NOT NULL,
            unit_serial TEXT NOT NULL,
            lab_number TEXT,
            chps_number TEXT,
            address_code TEXT,
            cert_date TEXT,
            shipped_date TEXT,
            antenna1_number TEXT,
            antenna2_number TEXT
        )
    """)
    conn.execute("CREATE INDEX IF NOT EXISTS idx_records_sheet_serial ON records(sheet_name, unit_serial)")
    conn.execute("CREATE INDEX IF NOT EXISTS idx_records_source ON records(source_path)")
    conn.commit()
    return conn

def _clean_history_row(row_tuple, sheet_name, source_file, source_path, file_order, row_order):
    if not row_tuple or len(row_tuple) < 5:
        return None
    cell_b = row_tuple[1]
    if not cell_b:
        return None

    is_radar = (sheet_name == 'RADAR')
    chps_val = row_tuple[2]
    ant1_val = row_tuple[11] if is_radar and len(row_tuple) > 11 else None
    ant2_val = row_tuple[12] if is_radar and len(row_tuple) > 12 else None
    chps_clean = str(chps_val).replace('CHPS', '') if chps_val else ''
    ant1_clean = ant1_val.replace('S/N ', '') if isinstance(ant1_val, str) and ant1_val != 'N/A' else None
    ant2_clean = ant2_val.replace('S/N ', '') if isinstance(ant2_val, str) and ant2_val != 'N/A' else None
    ship_val = row_tuple[5] if len(row_tuple) > 5 else None

    return {
        'source_path':     source_path,
        'source_file':     source_file,
        'file_order':      file_order,
        'row_order':       row_order,
        'sheet_name':      sheet_name,
        'lab_number':      row_tuple[0],
        'unit_serial':     str(cell_b).strip(),
        'chps_number':     chps_clean,
        'address_code':    normalize_address_code(row_tuple[3]),
        'date':            row_tuple[4],
        'shipped_date':    ship_val,
        'antenna1_number': ant1_clean,
        'antenna2_number': ant2_clean,
    }

def _db_row_to_record(row):
    return {
        'lab_number':      row['lab_number'],
        'unit_serial':     row['unit_serial'],
        'chps_number':     row['chps_number'] or '',
        'address_code':    normalize_address_code(row['address_code']),
        'date':            row['cert_date'],
        'shipped_date':    row['shipped_date'],
        'antenna1_number': row['antenna1_number'],
        'antenna2_number': row['antenna2_number'],
        'source_file':     row['source_file'],
        'source_path':     row['source_path'],
        'file_order':      row['file_order'],
        'row_order':       row['row_order'],
    }

def _insert_history_record(conn, record):
    conn.execute("""
        INSERT INTO records (
            source_path, source_file, file_order, row_order, sheet_name,
            unit_serial, lab_number, chps_number, address_code, cert_date,
            shipped_date, antenna1_number, antenna2_number
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        str(record['source_path']),
        str(record['source_file']),
        int(record['file_order']),
        int(record['row_order']),
        str(record['sheet_name']),
        str(record['unit_serial']),
        str(record['lab_number']) if record['lab_number'] is not None else None,
        str(record['chps_number']) if record['chps_number'] is not None else None,
        str(record['address_code']) if record['address_code'] is not None else None,
        str(record['date']) if record['date'] is not None else None,
        str(record['shipped_date']) if record['shipped_date'] is not None else None,
        str(record['antenna1_number']) if record['antenna1_number'] is not None else None,
        str(record['antenna2_number']) if record['antenna2_number'] is not None else None,
    ))

def _index_workbook_to_db(conn, filepath, file_order):
    fname = os.path.basename(filepath)
    mtime_ns, size = _file_stamp(filepath)

    conn.execute("DELETE FROM records WHERE source_path = ?", (filepath,))
    wb = load_workbook(filepath, read_only=True, data_only=True)
    try:
        for sheet_name in ('RADAR', 'LIDAR'):
            if sheet_name not in wb.sheetnames:
                continue
            ws = wb[sheet_name]
            for row_order, row_tuple in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
                record = _clean_history_row(row_tuple, sheet_name, fname, filepath, file_order, row_order)
                if record:
                    _insert_history_record(conn, record)
    finally:
        wb.close()

    conn.execute("""
        INSERT OR REPLACE INTO files (path, mtime_ns, size, file_order, indexed_at)
        VALUES (?, ?, ?, ?, ?)
    """, (filepath, mtime_ns, size, file_order, datetime.now().isoformat(timespec='seconds')))

def _load_history_index_from_db(conn):
    conn.row_factory = sqlite3.Row
    index = {'RADAR': {}, 'LIDAR': {}}
    for row in conn.execute("""
        SELECT * FROM records
        ORDER BY file_order ASC, row_order ASC
    """):
        record = _db_row_to_record(row)
        index[row['sheet_name']][row['unit_serial']] = record
    return index

def _search_history_db(serial_number, sheet_name, all_matches=False):
    try:
        conn = _history_db_connect()
        conn.row_factory = sqlite3.Row
        rows = []
        seen_ids = set()
        for term in serial_search_terms(serial_number):
            for row in conn.execute("""
                SELECT * FROM records
                WHERE sheet_name = ? AND unit_serial LIKE ?
                ORDER BY file_order DESC, row_order DESC
            """, (sheet_name, f"%{term}")).fetchall():
                row_id = row['id']
                if row_id not in seen_ids:
                    seen_ids.add(row_id)
                    rows.append(row)
        conn.close()
        records = [_db_row_to_record(row) for row in rows]
        return records if all_matches else choose_best_history_record(records)
    except Exception:
        return [] if all_matches else None

def update_history_index_record(sheet_name, record):
    """Keep the runtime history index current after saving a new Excel row."""
    try:
        _HISTORY_INDEX.setdefault(sheet_name, {})[record['unit_serial']] = record
    except Exception:
        pass

def _build_history_index(on_progress=None):
    """
    Build the in-memory index from a persistent SQLite cache.
    Excel files are only re-read when their mtime/size changes.
    """
    global _HISTORY_INDEX, _INDEX_READY, _INDEX_STATUS

    _INDEX_READY  = False
    _INDEX_STATUS = "Checking index cache..."

    try:
        conn = _history_db_connect()
        all_files = [p for p in _history_workbooks() if os.path.exists(p)]
        total = len(all_files)
        current_paths = set(all_files)

        # Drop records for files that no longer exist.
        for (cached_path,) in conn.execute("SELECT path FROM files").fetchall():
            if cached_path not in current_paths:
                conn.execute("DELETE FROM records WHERE source_path = ?", (cached_path,))
                conn.execute("DELETE FROM files WHERE path = ?", (cached_path,))

        for i, filepath in enumerate(all_files):
            fname = os.path.basename(filepath)
            file_order = i
            try:
                mtime_ns, size = _file_stamp(filepath)
                cached = conn.execute(
                    "SELECT mtime_ns, size, file_order FROM files WHERE path = ?",
                    (filepath,)
                ).fetchone()
                needs_index = (
                    cached is None or
                    cached[0] != mtime_ns or
                    cached[1] != size or
                    cached[2] != file_order
                )
                if needs_index:
                    if on_progress:
                        on_progress(f"Indexing {i+1}/{total}: {fname}")
                    _index_workbook_to_db(conn, filepath, file_order)
                    conn.commit()
                else:
                    if on_progress:
                        on_progress(f"Using cache {i+1}/{total}: {fname}")
            except Exception as e:
                if on_progress:
                    on_progress(f"Skipped {fname}: {e}")
                continue

        index = _load_history_index_from_db(conn)
        conn.close()

        _HISTORY_INDEX = index
        _INDEX_READY   = True
        radar_count = len(index['RADAR'])
        lidar_count = len(index['LIDAR'])
        _INDEX_STATUS  = f"Index ready — {radar_count} RADAR, {lidar_count} LIDAR records"
        if on_progress:
            on_progress(_INDEX_STATUS)
    except Exception as e:
        _HISTORY_INDEX = {'RADAR': {}, 'LIDAR': {}}
        _INDEX_READY   = False
        _INDEX_STATUS  = f"Index cache unavailable: {e}"
        if on_progress:
            on_progress(_INDEX_STATUS)


def search_all_history(serial_number, sheet_name):
    """
    Search the in-memory index for the most recent entry matching serial_number.
    Falls back to SQLite cache, then file scan if index not ready yet.
    """
    if _INDEX_READY:
        sheet_index = _HISTORY_INDEX.get(sheet_name, {})
        matches = [
            record for key, record in sheet_index.items()
            if any(str(key).endswith(term) for term in serial_search_terms(serial_number))
        ]
        return choose_best_history_record(matches)

    cached_record = _search_history_db(serial_number, sheet_name, all_matches=False)
    if cached_record:
        return cached_record

    # Fallback: original file scan
    all_files = [EXCEL_2026] + sorted(glob.glob(os.path.join(HISTORY_DIR, 'week_report_*.xlsx')), reverse=True)

    for filepath in all_files:
        try:
            wb = load_workbook(filepath, read_only=True, data_only=True)
            if sheet_name not in wb.sheetnames:
                wb.close()
                continue
            ws = wb[sheet_name]
            for row in range(ws.max_row, 1, -1):
                cell = ws.cell(row=row, column=2).value
                if cell and any(str(cell).endswith(term) for term in serial_search_terms(serial_number)):
                    date_val  = ws.cell(row=row, column=5).value
                    lab_val   = ws.cell(row=row, column=1).value
                    chps_val  = ws.cell(row=row, column=3).value
                    addr_val  = ws.cell(row=row, column=4).value
                    ant1_val  = ws.cell(row=row, column=12).value if sheet_name == 'RADAR' else None
                    ant2_val  = ws.cell(row=row, column=13).value if sheet_name == 'RADAR' else None
                    wb.close()
                    chps_clean = str(chps_val).replace('CHPS', '') if chps_val else ''
                    ant1_clean = ant1_val.replace('S/N ', '') if ant1_val and ant1_val != 'N/A' else None
                    ant2_clean = ant2_val.replace('S/N ', '') if ant2_val and ant2_val != 'N/A' else None
                    ship_val = ws.cell(row=row, column=6).value
                    return {
                        'lab_number':   lab_val,
                        'unit_serial':  cell,
                        'chps_number':  chps_clean,
                        'address_code': normalize_address_code(addr_val),
                        'date':         date_val,
                        'shipped_date': ship_val,
                        'antenna1_number': ant1_clean,
                        'antenna2_number': ant2_clean,
                        'source_file':  os.path.basename(filepath),
                        'source_path':  filepath,
                    }
            wb.close()
        except Exception:
            continue
    return None


# FIX 6: search_all_entries_for_serial — fetch ALL entries across ALL years
def search_all_entries_for_serial(serial_number, sheet_name):
    """
    Return a LIST of ALL records matching serial_number across all history files
    and the current year Excel. Newest entries appear first.
    Used by the Search Tab to show complete history.
    """
    results = []

    cached_results = _search_history_db(serial_number, sheet_name, all_matches=True)
    if cached_results:
        return cached_results

    all_files = [EXCEL_2026] + sorted(glob.glob(os.path.join(HISTORY_DIR, 'week_report_*.xlsx')), reverse=True)

    for filepath in all_files:
        fname = os.path.basename(filepath)
        try:
            wb = load_workbook(filepath, read_only=True, data_only=True)
            if sheet_name not in wb.sheetnames:
                wb.close()
                continue
            ws = wb[sheet_name]
            for row_tuple in ws.iter_rows(min_row=2, values_only=True):
                if not row_tuple or len(row_tuple) < 5:
                    continue
                cell_b = row_tuple[1]
                if not cell_b:
                    continue
                key = str(cell_b).strip()
                if any(key.endswith(term) for term in serial_search_terms(serial_number)):
                    chps_val = row_tuple[2]
                    ant1_val = row_tuple[11] if sheet_name == 'RADAR' and len(row_tuple) > 11 else None
                    ant2_val = row_tuple[12] if sheet_name == 'RADAR' and len(row_tuple) > 12 else None
                    chps_clean = str(chps_val).replace('CHPS', '') if chps_val else ''
                    ant1_clean = ant1_val.replace('S/N ', '') if isinstance(ant1_val, str) and ant1_val != 'N/A' else None
                    ant2_clean = ant2_val.replace('S/N ', '') if isinstance(ant2_val, str) and ant2_val != 'N/A' else None
                    ship_val = row_tuple[5] if len(row_tuple) > 5 else None
                    results.append({
                        'lab_number':      row_tuple[0],
                        'unit_serial':     key,
                        'chps_number':     chps_clean,
                        'address_code':    normalize_address_code(row_tuple[3]),
                        'date':            row_tuple[4],
                        'shipped_date':    ship_val,
                        'antenna1_number': ant1_clean,
                        'antenna2_number': ant2_clean,
                        'source_file':     fname,
                        'source_path':     filepath,
                    })
            wb.close()
        except Exception:
            continue
    return results


def years_since(date_val):
    """Return float years since a date value (string or datetime)."""
    if not date_val:
        return None
    if isinstance(date_val, str):
        for fmt in ("%m/%d/%Y", "%Y-%m-%d", "%Y-%m-%d %H:%M:%S", "%m/%d/%y"):
            try:
                date_val = datetime.strptime(date_val, fmt)
                break
            except ValueError:
                continue
        else:
            return None
    if hasattr(date_val, 'year'):
        delta = datetime.now() - date_val
        return delta.days / 365.25
    return None

def digits_match_transposed(a, b):
    """Check if b looks like a with two adjacent digits swapped."""
    if len(a) != len(b):
        return False
    diffs = [i for i in range(len(a)) if a[i] != b[i]]
    if len(diffs) == 2:
        i, j = diffs
        if j == i + 1 and a[i] == b[j] and a[j] == b[i]:
            return True
    return False

def parse_lab_number(value, prefix):
    """Return the numeric lab number from a full lab value or bare number."""
    raw = str(value or '').replace(' (RETEST)', '').strip()
    if not raw:
        return None
    upper_raw = raw.upper()
    upper_prefix = prefix.upper()
    if upper_raw.startswith(upper_prefix):
        raw = raw[len(prefix):].strip()
    if raw.isdigit():
        return int(raw)
    return None

def format_lab_number(prefix, number):
    return f"{prefix}{int(number)}"

def expand_log_number_entries(raw):
    """
    Expand comma/newline separated lab numbers plus ranges.
    Supported range examples: AS26-72 to AS26-80, AS26-72-AS26-80, ASL26-45:48.
    """
    import re

    entries = []
    errors = []
    chunks = [x.strip() for x in re.split(r'[,\n]+', raw or '') if x.strip()]
    single_re = re.compile(r'^(ASL26-|AS26-)(\d+)$', re.IGNORECASE)
    range_re = re.compile(
        r'^(ASL26-|AS26-)(\d+)\s*(?:TO|:|-)\s*(?:(ASL26-|AS26-)?)(\d+)$',
        re.IGNORECASE
    )

    for chunk in chunks:
        normalized = chunk.strip().upper()
        single = single_re.match(normalized.replace(' ', ''))
        if single:
            prefix, number = single.groups()
            entries.append(format_lab_number(prefix, number))
            continue

        match = range_re.match(normalized)
        if match:
            start_prefix, start_num, end_prefix, end_num = match.groups()
            end_prefix = end_prefix or start_prefix
            if start_prefix.upper() != end_prefix.upper():
                errors.append(f"{chunk}: range prefixes must match")
                continue
            start = int(start_num)
            end = int(end_num)
            if end < start:
                errors.append(f"{chunk}: range end is before start")
                continue
            for number in range(start, end + 1):
                entries.append(format_lab_number(start_prefix.upper(), number))
            continue

        errors.append(f"{chunk}: expected AS26-### or ASL26-###")

    unique_entries = []
    seen = set()
    for entry in entries:
        key = entry.upper()
        if key not in seen:
            seen.add(key)
            unique_entries.append(entry)
    return unique_entries, errors

def normalize_scanned_value(value, field_label="", expected_length=None, any_length=False):
    """
    Normalize common keyboard-wedge barcode scans before verification.
    Handles prefixes such as CHPS, S/N, SN, serial labels, and unit type prefixes.
    """
    import re

    raw = str(value or '').strip()
    if not raw:
        return ''

    # Scanners often append Enter/Tab and sometimes wrap data in whitespace.
    raw = raw.replace('\r', '').replace('\n', '').replace('\t', '').strip()
    label = (field_label or '').upper()
    upper_raw = raw.upper()

    if 'CHPS' in label:
        m = re.search(r'CHPS\s*[-:#]?\s*(\d+)', upper_raw)
        if m:
            return m.group(1)[-expected_length:] if expected_length else m.group(1)
        digits = ''.join(re.findall(r'\d+', raw))
        if expected_length and len(digits) >= expected_length:
            return digits[-expected_length:]
        return digits or raw

    if any(token in label for token in ('SERIAL', 'ANTENNA', 'FA NUMBER', 'FB NUMBER')):
        cleaned = re.sub(r'(?i)\b(SERIAL|SERIAL\s*NUMBER|S/N|SN|SNO|NO|NUMBER)\b', '', raw)
        cleaned = cleaned.strip(' :-#')
        for unit_prefix in sorted(RADAR_TYPES + LIDAR_TYPES, key=len, reverse=True):
            if cleaned.upper().startswith(unit_prefix) and len(cleaned) > len(unit_prefix):
                cleaned = cleaned[len(unit_prefix):].strip(' :-#')
                break
        if any_length:
            return cleaned or raw
        digits = ''.join(re.findall(r'\d+', cleaned))
        if expected_length and not any_length and len(digits) >= expected_length:
            return digits[-expected_length:]
        if digits:
            return digits
        return cleaned or raw

    return raw

def normalize_address_code(value):
    """Accept scans like 'CHP (530)' or 'Address: 530' and return the code."""
    import re

    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    if isinstance(value, int):
        return str(value)

    raw = str(value or '').strip()
    if not raw:
        return ''
    m = re.search(r'CHP\s*\(([^)]+)\)', raw, flags=re.IGNORECASE)
    if m:
        return m.group(1).strip()
    decimal_m = re.fullmatch(r'(\d+)\.0+', raw)
    if decimal_m:
        return decimal_m.group(1)
    digits = ''.join(re.findall(r'\d+', raw))
    if digits:
        return digits[-3:] if len(digits) > 3 else digits
    return raw

def detect_unit_type_prefix(value):
    """Return a known RADAR/LIDAR unit type prefix from a scanned unit value."""
    raw = str(value or '').strip().upper()
    for unit_prefix in sorted(RADAR_TYPES + LIDAR_TYPES, key=len, reverse=True):
        if raw.startswith(unit_prefix) and len(raw) > len(unit_prefix):
            return unit_prefix
    return None

def is_blank_excel_value(value):
    if value is None:
        return True
    return str(value).strip() in ('', ' ', 'None', '—')

def is_current_year_record(history):
    if not history:
        return False
    current_file = os.path.basename(EXCEL_2026)
    return (
        history.get('source_path') == EXCEL_2026 or
        history.get('source_file') == current_file
    )

def is_current_year_failed_record(history):
    return is_current_year_record(history) and is_blank_excel_value(history.get('shipped_date'))

def serial_search_terms(serial_number):
    raw = str(serial_number or '').strip()
    terms = []
    if raw:
        terms.append(raw)
        stripped = raw.lstrip('0')
        if stripped and stripped != raw:
            terms.append(stripped)
    return terms

def choose_best_history_record(records):
    """Prefer same-year failed/current-year records over older matching serials."""
    if not records:
        return None

    def score(record):
        file_order = record.get('file_order')
        row_order = record.get('row_order')
        try:
            file_order = int(file_order)
        except (TypeError, ValueError):
            file_order = -1
        try:
            row_order = int(row_order)
        except (TypeError, ValueError):
            row_order = -1
        return (
            1 if is_current_year_failed_record(record) else 0,
            1 if is_current_year_record(record) else 0,
            file_order,
            row_order,
        )

    return max(records, key=score)


# ─────────────────────────────────────────────
#  VERIFY DIALOG
# ─────────────────────────────────────────────
class VerifyDialog(tk.Toplevel):
    """
    Single-entry + digit readback dialog.
    Returns confirmed value via self.result (None if cancelled).
    """
    def __init__(self, parent, field_label, prefill="", length=6, any_length=False):
        super().__init__(parent)
        self.result = None
        self.field_label = field_label
        self.required_length = length
        self.any_length = any_length

        self.title(f"Verify: {field_label}")
        self.configure(bg=WHITE)
        self.resizable(False, False)
        self.grab_set()
        self.focus_set()

        self._phase = "entry"
        self._prefill = prefill
        if self._prefill:
            if not self._check_single_value(self._prefill):
                self._build_entry_phase(self._prefill)
        else:
            self._build_entry_phase()

        self.update_idletasks()
        px = parent.winfo_rootx() + (parent.winfo_width()  - self.winfo_width())  // 2
        py = parent.winfo_rooty() + (parent.winfo_height() - self.winfo_height()) // 2
        self.geometry(f"+{px}+{py}")

    def _clear_frame(self):
        for w in self.winfo_children():
            w.destroy()

    def _build_entry_phase(self, initial_value=""):
        self._clear_frame()
        tk.Label(self, text=f"VERIFY: {self.field_label}",
                 bg=WHITE, fg=ACCENT,
                 font=("Courier New", 12, "bold")).pack(padx=20, pady=(20,4))

        tk.Label(self, text="Enter value once:", bg=WHITE, fg=TEXT,
                 font=("Courier New", 10)).pack(anchor='w', padx=20)
        self.entry1 = tk.Entry(self, bg=PANEL_BG, fg=TEXT, insertbackground=TEXT,
                               font=("Courier New", 16, "bold"),
                               width=18, relief='solid', bd=1)
        self.entry1.pack(padx=20, pady=4, ipady=6)
        if initial_value:
            self.entry1.insert(0, initial_value)

        self.msg_var = tk.StringVar()
        tk.Label(self, textvariable=self.msg_var, bg=WHITE, fg=RED,
                 font=("Courier New", 9)).pack(pady=2)

        btn_frame = tk.Frame(self, bg=WHITE)
        btn_frame.pack(pady=(4, 20))
        tk.Button(btn_frame, text="CHECK", bg=ACCENT, fg=WHITE,
                  font=("Courier New", 10, "bold"), relief='flat',
                  padx=16, pady=6, cursor='hand2',
                  command=self._check_match).pack(side='left', padx=6)
        tk.Button(btn_frame, text="CANCEL", bg=PANEL_BG, fg=TEXT,
                  font=("Courier New", 10), relief='flat',
                  padx=16, pady=6, cursor='hand2',
                  command=self.destroy).pack(side='left', padx=6)

        self.entry1.bind("<Return>", lambda e: self._check_match())
        self.entry1.focus_set()

    def _check_match(self):
        v1 = self.entry1.get().strip()
        self._check_single_value(v1)

    def _check_single_value(self, value):
        if not value:
            if hasattr(self, 'msg_var'):
                self.msg_var.set("⚠  Value required.")
            return False

        if not self.any_length and len(value) != self.required_length:
            if hasattr(self, 'msg_var'):
                self.msg_var.set(f"⚠  Must be {self.required_length} characters (got {len(value)}).")
            return False

        self._confirmed_value = value
        self._build_readback_phase(value)
        return True

    def _build_readback_phase(self, value):
        self._clear_frame()

        tk.Label(self, text=f"CONFIRM: {self.field_label}",
                 bg=WHITE, fg=ACCENT,
                 font=("Courier New", 12, "bold")).pack(padx=20, pady=(20, 4))

        tk.Label(self, text="Read each digit carefully:",
                 bg=WHITE, fg=TEXT_DIM,
                 font=("Courier New", 9)).pack()

        digit_frame = tk.Frame(self, bg=WHITE)
        digit_frame.pack(padx=20, pady=12)

        positions = ['1ST','2ND','3RD','4TH','5TH','6TH','7TH','8TH','9TH','10TH']
        for i, ch in enumerate(value):
            box = tk.Frame(digit_frame, bg=PANEL_BG, bd=1, relief='solid',
                           width=54, height=60)
            box.pack_propagate(False)
            box.pack(side='left', padx=3)
            tk.Label(box, text=ch, bg=PANEL_BG, fg=ACCENT,
                     font=("Courier New", 22, "bold")).pack(expand=True)
            pos = positions[i] if i < len(positions) else f"{i+1}TH"
            tk.Label(box, text=pos, bg=PANEL_BG, fg=TEXT_DIM,
                     font=("Courier New", 7)).pack(pady=(0, 4))

        tk.Label(self, text=f"Full value:  {value}",
                 bg=WHITE, fg=TEXT,
                 font=("Courier New", 13, "bold")).pack(pady=6)

        btn_frame = tk.Frame(self, bg=WHITE)
        btn_frame.pack(pady=(4, 20))
        confirm_btn = tk.Button(btn_frame, text="✓  CONFIRM — CORRECT",
                  bg=GREEN, fg=WHITE,
                  font=("Courier New", 10, "bold"), relief='flat',
                  padx=14, pady=7, cursor='hand2',
                  command=self._confirm)
        confirm_btn.pack(side='left', padx=6)
        confirm_btn.focus_set()
        confirm_btn.bind("<Return>", lambda e: self._confirm())
        reenter_btn = tk.Button(btn_frame, text="✗  RE-ENTER",
                  bg=RED, fg=WHITE,
                  font=("Courier New", 10), relief='flat',
                  padx=14, pady=7, cursor='hand2',
                  command=self._build_entry_phase)
        reenter_btn.pack(side='left', padx=6)
        reenter_btn.bind("<Return>", lambda e: self._build_entry_phase())
        confirm_btn.bind("<Tab>", lambda e: (reenter_btn.focus_set(), "break"))
        reenter_btn.bind("<Tab>", lambda e: (confirm_btn.focus_set(), "break"))

    def _confirm(self):
        self.result = self._confirmed_value
        self.destroy()


# ─────────────────────────────────────────────
#  SEAL PLACEMENT
# ─────────────────────────────────────────────
def find_cert_row(doc, label_text):
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            for cell in row.cells:
                txt = cell.text.strip()
                if txt.startswith(label_text):
                    return table, i
    return None, None

def find_cert_row_by_label(doc, labels):
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                txt = cell.text.strip()
                for lbl in labels:
                    if txt == lbl or txt.upper() == lbl.upper():
                        return table, i, j
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                txt = cell.text.strip().upper()
                for lbl in labels:
                    if lbl.upper() in txt and len(txt) < len(lbl) + 15:
                        return table, i, j
    return None, None, None

def add_floating_image(doc, image_path, target_para, width_cm=4.0, height_cm=4.0,
                        x_cm=0.0, y_cm=0.0):
    from docx.shared import Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import lxml.etree as etree

    EMU = 914400 / 2.54
    w_emu = int(width_cm  * EMU)
    h_emu = int(height_cm * EMU)
    x_emu = int(x_cm      * EMU)
    y_emu = int(y_cm      * EMU)
    import random
    unique_id = random.randint(100, 9999)

    run = target_para.add_run()
    inline_pic = run.add_picture(image_path, width=Cm(width_cm), height=Cm(height_cm))

    inline_el = run._r.find(qn('w:drawing'))
    if inline_el is None:
        return

    inline_body = inline_el.find(qn('wp:inline'))
    if inline_body is None:
        return

    graphic = inline_body.find(qn('a:graphic'))
    extent  = inline_body.find(qn('wp:extent'))

    anchor_xml = f"""<wp:anchor
        xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
        xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"
        xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        distT="0" distB="0" distL="0" distR="0"
        simplePos="0" relativeHeight="251658240" behindDoc="0"
        locked="0" layoutInCell="1" allowOverlap="1">
      <wp:simplePos x="0" y="0"/>
      <wp:positionH relativeFrom="column">
        <wp:posOffset>{x_emu}</wp:posOffset>
      </wp:positionH>
      <wp:positionV relativeFrom="paragraph">
        <wp:posOffset>{y_emu}</wp:posOffset>
      </wp:positionV>
      <wp:extent cx="{w_emu}" cy="{h_emu}"/>
      <wp:effectExtent l="0" t="0" r="0" b="0"/>
      <wp:wrapNone/>
      <wp:docPr id="{unique_id}" name="Seal"/>
      <wp:cNvGraphicFramePr>
        <a:graphicFrameLocks xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
                             noChangeAspect="1"/>
      </wp:cNvGraphicFramePr>
      {etree.tostring(graphic, encoding="unicode")}
    </wp:anchor>"""

    anchor_el = etree.fromstring(anchor_xml)
    inline_el.remove(inline_body)
    inline_el.append(anchor_el)

def seal_document(src_path, dst_path, unit_type, cert_date=""):
    import shutil

    if not os.path.exists(SEAL_IMAGE):
        return False, f"Seal image not found:\n{SEAL_IMAGE}"

    if not os.path.exists(src_path):
        return False, f"Certificate not found:\n{src_path}"

    os.makedirs(SEALED_DIR, exist_ok=True)
    shutil.copy2(src_path, dst_path)

    try:
        doc = docx.Document(dst_path)
    except Exception as e:
        return False, f"Cannot open certificate:\n{e}"

    search_labels = ["CERTIFICATION"]

    table, row_idx, label_col = find_cert_row_by_label(doc, search_labels)

    if table is not None and row_idx is not None:
        row   = table.rows[row_idx]

        seen, unique = set(), []
        for c in row.cells:
            if id(c) not in seen:
                seen.add(id(c))
                unique.append(c)

        target_cell = None
        passed = False
        for c in unique:
            if any(lbl.upper() in c.text.strip().upper() for lbl in search_labels):
                passed = True
                continue
            if passed:
                target_cell = c
                break
        if target_cell is None and len(unique) > 1:
            target_cell = unique[-1]

        if target_cell is None:
            return False, "CERTIFICATION row found but no target cell."

        if target_cell.paragraphs:
            anchor_para = target_cell.paragraphs[0]
        else:
            anchor_para = target_cell.add_paragraph()

        try:
            if unit_type == "RADAR":
                seal_x, seal_y = -4.8, -1.8
            else:
                seal_x, seal_y = -4.8, -1.2

            add_floating_image(
                doc, SEAL_IMAGE,
                target_para=anchor_para,
                width_cm=5.5, height_cm=5.5,
                x_cm=seal_x, y_cm=seal_y
            )

            if cert_date:
                add_floating_text(
                    doc, cert_date,
                    target_para=anchor_para,
                    font_size_pt=11,
                    bold=True,
                    color="FF0000",
                    x_cm=seal_x + 2.0,
                    y_cm=seal_y + 2.6,
                    width_cm=4.0,
                    height_cm=0.7
                )

            if cert_date:
                fill_date_certified(doc, cert_date)

            doc.save(dst_path)
            return True, ""
        except Exception as e:
            return False, f"Failed to place floating seal:\n{e}"
    else:
        return False, (
            "Could not find CERTIFICATION row in document tables.\n"
            "Run diagnose_cert.py to inspect the cert structure."
        )


def add_floating_text(doc, text, target_para, font_size_pt=10,
                       x_cm=0.0, y_cm=0.0, width_cm=5.0, height_cm=1.0,
                       bold=False, color="000000"):
    from docx.shared import Cm, Pt
    from docx.oxml.ns import qn
    import lxml.etree as etree

    EMU = 360000
    x_emu = int(x_cm     * EMU)
    y_emu = int(y_cm     * EMU)
    w_emu = int(width_cm * EMU)
    h_emu = int(height_cm* EMU)
    pt_emu = int(font_size_pt * 12700)
    bold_val = "1" if bold else "0"

    import random
    unique_id = random.randint(100, 9999)
    txbx_xml = f"""<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:rPr/>
  <mc:AlternateContent xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006">
    <mc:Choice Requires="wps">
      <w:drawing>
        <wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                   distT="0" distB="0" distL="0" distR="0"
                   simplePos="0" relativeHeight="251658241" behindDoc="0"
                   locked="0" layoutInCell="1" allowOverlap="1">
          <wp:simplePos x="0" y="0"/>
          <wp:positionH relativeFrom="column">
            <wp:posOffset>{x_emu}</wp:posOffset>
          </wp:positionH>
          <wp:positionV relativeFrom="paragraph">
            <wp:posOffset>{y_emu}</wp:posOffset>
          </wp:positionV>
          <wp:extent cx="{w_emu}" cy="{h_emu}"/>
          <wp:effectExtent l="0" t="0" r="0" b="0"/>
          <wp:wrapNone/>
          <wp:docPr id="{unique_id}" name="DateText"/>
          <wp:cNvGraphicFramePr/>
          <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
            <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
              <wps:wsp xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
                <wps:cNvSpPr><a:spLocks noChangeArrowheads="1"/></wps:cNvSpPr>
                <wps:spPr>
                  <a:xfrm><a:off x="0" y="0"/><a:ext cx="{w_emu}" cy="{h_emu}"/></a:xfrm>
                  <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
                  <a:noFill/>
                  <a:ln><a:noFill/></a:ln>
                </wps:spPr>
                <wps:txbx>
                  <w:txbxContent xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                    <w:p>
                      <w:pPr><w:jc w:val="left"/></w:pPr>
                      <w:r>
                        <w:rPr>
                          <w:rFonts w:ascii="Open Sans" w:hAnsi="Open Sans" w:cs="Open Sans"/>
                          <w:sz w:val="{font_size_pt * 2}"/>
                          <w:szCs w:val="{font_size_pt * 2}"/>
                          <w:b w:val="{bold_val}"/>
                          <w:color w:val="{color}"/>
                        </w:rPr>
                        <w:t xml:space="preserve">{text}</w:t>
                      </w:r>
                    </w:p>
                  </w:txbxContent>
                </wps:txbx>
                <wps:bodyPr rot="0" spcFirstLastPara="0" vertOverflow="overflow"
                            horzOverflow="overflow" vert="horz" wrap="none"
                            lIns="0" tIns="0" rIns="0" bIns="0" numCol="1"
                            spcCol="0" rtlCol="0" fromWordArt="0" anchor="t"
                            anchorCtr="0" forceAA="0" compatLnSpc="1">
                  <a:spAutoFit/>
                </wps:bodyPr>
              </wps:wsp>
            </a:graphicData>
          </a:graphic>
        </wp:anchor>
      </w:drawing>
    </mc:Choice>
  </mc:AlternateContent>
</w:r>"""

    txbx_el = etree.fromstring(txbx_xml)
    target_para._p.append(txbx_el)


def fill_date_certified(doc, cert_date):
    from docx.oxml.ns import qn
    import lxml.etree as etree

    def style_run(run):
        run.bold = True
        rPr = run._r.get_or_add_rPr()
        for tag in [qn("w:rFonts"), qn("w:color")]:
            for el in rPr.findall(tag):
                rPr.remove(el)
        rFonts = etree.SubElement(rPr, qn("w:rFonts"))
        rFonts.set(qn("w:ascii"), "Arial")
        rFonts.set(qn("w:hAnsi"), "Arial")
        rFonts.set(qn("w:cs"),    "Arial")
        color_el = etree.SubElement(rPr, qn("w:color"))
        color_el.set(qn("w:val"), "000000")

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                stripped = cell.text.strip()
                if stripped.lower() == "date certified:":
                    para = cell.paragraphs[0]
                    run0 = para.runs[0] if para.runs else para.add_run()
                    run0.text = "Date certified:  " + cert_date
                    style_run(run0)
                    return True, ""

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip().lower() == "date certified:":
                        run0 = para.runs[0] if para.runs else para.add_run()
                        run0.text = "Date certified:  " + cert_date
                        style_run(run0)
                        return True, ""

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if "date certified:" in para.text.lower():
                        for run in reversed(para.runs):
                            if "date certified" in run.text.lower():
                                import re
                                run.text = re.sub(
                                    r'(Date certified:\s*).*',
                                    lambda m, d=cert_date: m.group(1) + "  " + d,
                                    run.text,
                                    flags=re.IGNORECASE
                                )
                                style_run(run)
                                return True, ""

    return False, "'Date certified:' not found anywhere in document"

# ─────────────────────────────────────────────
#  STARTUP FILE LOCK CHECK
# ─────────────────────────────────────────────
def get_files_to_check():
    files = [
        ("Excel Log (week_report_2026.xlsx)", EXCEL_2026),
        ("Area Code Doc (area_code.docx)",    AREA_CODE_FILE),
    ]
    for name, fname in RADAR_TEMPLATES.items():
        files.append((f"RADAR Template ({fname})", os.path.join(TEMPLATE_DIR, fname)))
    for name, fname in LIDAR_TEMPLATES.items():
        files.append((f"LIDAR Template ({fname})", os.path.join(TEMPLATE_DIR, fname)))
    return files

def check_file_locked(path):
    if not os.path.exists(path):
        return False
    try:
        with open(path, 'a+b'):
            pass
        return False
    except (IOError, PermissionError):
        return True

def check_all_files_free():
    locked = []
    for label, path in get_files_to_check():
        if check_file_locked(path):
            locked.append((label, path))
    return locked

def startup_file_check(root):
    while True:
        locked = check_all_files_free()
        if not locked:
            return True

        file_list = "\n".join(f"  •  {label}" for label, _ in locked)
        msg = (
            f"The following file(s) are currently open in Word or Excel:\n\n"
            f"{file_list}\n\n"
            f"Please close them and click Retry, or click Exit to quit."
        )
        dlg = tk.Toplevel(root)
        dlg.title("Files In Use — Please Close")
        dlg.configure(bg=WHITE)
        dlg.resizable(False, False)
        dlg.grab_set()

        dlg.update_idletasks()
        sw = dlg.winfo_screenwidth()
        sh = dlg.winfo_screenheight()
        dlg.geometry(f"520x{min(300 + len(locked)*22, 500)}+{(sw-520)//2}+{(sh-400)//2}")

        tk.Label(dlg, text="⚠  FILES IN USE", bg=WHITE, fg=RED,
                 font=("Courier New", 13, "bold")).pack(pady=(20, 4))
        tk.Label(dlg,
                 text="Close these files before the app can start:",
                 bg=WHITE, fg=TEXT,
                 font=("Courier New", 10)).pack(pady=(0, 8))

        list_frame = tk.Frame(dlg, bg=PANEL_BG, relief='solid', bd=1)
        list_frame.pack(fill='x', padx=24, pady=4)
        for label, path in locked:
            row = tk.Frame(list_frame, bg=PANEL_BG)
            row.pack(fill='x', padx=8, pady=3)
            tk.Label(row, text="●", bg=PANEL_BG, fg=RED,
                     font=("Courier New", 10)).pack(side='left', padx=(0,6))
            tk.Label(row, text=label, bg=PANEL_BG, fg=TEXT,
                     font=("Courier New", 10), anchor='w').pack(side='left')

        tk.Label(dlg,
                 text="Close the file(s) above, then click Retry.",
                 bg=WHITE, fg=TEXT_DIM,
                 font=("Courier New", 9)).pack(pady=(10, 4))

        result = [None]

        btn_frame = tk.Frame(dlg, bg=WHITE)
        btn_frame.pack(pady=12)

        def on_retry():
            result[0] = 'retry'
            dlg.destroy()

        def on_exit():
            result[0] = 'exit'
            dlg.destroy()

        retry_btn = tk.Button(btn_frame, text="↺  RETRY",
                  bg=GREEN, fg=WHITE,
                  font=("Courier New", 10, "bold"), relief='flat',
                  padx=16, pady=8, cursor='hand2',
                  command=on_retry)
        retry_btn.pack(side='left', padx=8)
        retry_btn.focus_set()
        retry_btn.bind("<Return>", lambda e: on_retry())

        tk.Button(btn_frame, text="EXIT",
                  bg=PANEL_BG, fg=RED,
                  font=("Courier New", 10), relief='flat',
                  padx=16, pady=8, cursor='hand2',
                  command=on_exit).pack(side='left', padx=8)

        dlg.protocol("WM_DELETE_WINDOW", on_exit)
        root.wait_window(dlg)

        if result[0] == 'exit' or result[0] is None:
            return False


def wait_for_files_free(root, paths_with_labels):
    while True:
        locked = [(lbl, p) for lbl, p in paths_with_labels
                  if check_file_locked(p)]
        if not locked:
            return True

        file_list = "\n".join(f"  •  {lbl}" for lbl, _ in locked)
        dlg = tk.Toplevel(root)
        dlg.title("File In Use — Close and Retry")
        dlg.configure(bg=WHITE)
        dlg.resizable(False, False)
        dlg.grab_set()
        dlg.update_idletasks()
        sw = dlg.winfo_screenwidth()
        sh = dlg.winfo_screenheight()
        dlg.geometry(f"520x{min(280 + len(locked)*24, 480)}+{(sw-520)//2}+{(sh-400)//2}")

        tk.Label(dlg, text="⚠  CANNOT SAVE — FILE IN USE",
                 bg=WHITE, fg=RED,
                 font=("Courier New", 12, "bold")).pack(pady=(20, 4))
        tk.Label(dlg,
                 text="Your entry is safe. Close the file(s) below,\nthen click Retry to complete the save.",
                 bg=WHITE, fg=TEXT,
                 font=("Courier New", 10), justify='center').pack(pady=(0, 10))

        list_frame = tk.Frame(dlg, bg=PANEL_BG, relief='solid', bd=1)
        list_frame.pack(fill='x', padx=24, pady=4)
        for lbl, path in locked:
            row = tk.Frame(list_frame, bg=PANEL_BG)
            row.pack(fill='x', padx=8, pady=3)
            tk.Label(row, text="●", bg=PANEL_BG, fg=RED,
                     font=("Courier New", 10)).pack(side='left', padx=(0, 6))
            tk.Label(row, text=lbl, bg=PANEL_BG, fg=TEXT,
                     font=("Courier New", 10), anchor='w').pack(side='left')

        tk.Label(dlg, text="Your entered values are preserved.",
                 bg=WHITE, fg=GREEN,
                 font=("Courier New", 9, "bold")).pack(pady=(8, 4))

        result = [None]

        btn_frame = tk.Frame(dlg, bg=WHITE)
        btn_frame.pack(pady=12)

        def on_retry():
            result[0] = 'retry'
            dlg.destroy()

        def on_cancel():
            result[0] = 'cancel'
            dlg.destroy()

        retry_btn = tk.Button(btn_frame, text="↺  RETRY SAVE",
                  bg=GREEN, fg=WHITE,
                  font=("Courier New", 10, "bold"), relief='flat',
                  padx=16, pady=8, cursor='hand2',
                  command=on_retry)
        retry_btn.pack(side='left', padx=8)
        retry_btn.focus_set()
        retry_btn.bind("<Return>", lambda e: on_retry())

        tk.Button(btn_frame, text="CANCEL SAVE",
                  bg=PANEL_BG, fg=RED,
                  font=("Courier New", 10), relief='flat',
                  padx=16, pady=8, cursor='hand2',
                  command=on_cancel).pack(side='left', padx=8)

        dlg.protocol("WM_DELETE_WINDOW", on_cancel)
        root.wait_window(dlg)

        if result[0] != 'retry':
            return False

# ─────────────────────────────────────────────
#  SESSION PERSISTENCE
# ─────────────────────────────────────────────
def save_session(units):
    try:
        with open(SESSION_FILE, 'w', encoding='utf-8') as f:
            json.dump(units, f, indent=2, default=str)
    except Exception as e:
        print(f"[WARN] Could not save session: {e}")

def load_session():
    try:
        if os.path.exists(SESSION_FILE):
            with open(SESSION_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
    except Exception as e:
        print(f"[WARN] Could not load session: {e}")
    return []

def clear_session_file():
    try:
        if os.path.exists(SESSION_FILE):
            os.remove(SESSION_FILE)
    except Exception as e:
        print(f"[WARN] Could not clear session file: {e}")

# ─────────────────────────────────────────────
#  VERIFIED FIELD WIDGET
# ─────────────────────────────────────────────
class VerifiedField(tk.Frame):
    def __init__(self, parent, label, length=6, any_length=False,
                 prefill="", on_verified=None, next_field=None, **kwargs):
        super().__init__(parent, bg=PANEL_BG, **kwargs)
        self.label_text  = label
        self.length      = length
        self.any_length  = any_length
        self.on_verified = on_verified
        self.next_field  = next_field
        self._locked     = False
        self._value      = ""
        self._prefill    = prefill
        self.last_raw_value = ""

        self._build()
        if prefill:
            self._set_prefill(prefill)

    def _build(self):
        self.border_frame = tk.Frame(self, bg=RED, padx=2, pady=2)
        self.border_frame.grid(row=0, column=0, padx=(0, 6))

        self.display_var = tk.StringVar()
        self.entry_display = tk.Entry(
            self.border_frame,
            textvariable=self.display_var,
            bg=ENTRY_BG, fg=TEXT,
            insertbackground=TEXT,
            font=("Courier New", 13, "bold"),
            width=16, relief='flat', bd=4,
        )
        self.entry_display.pack(ipady=5)

        self.entry_display.bind("<Return>", lambda e: self._open_verify())
        self.entry_display.bind("<KP_Enter>", lambda e: self._open_verify())
        self.entry_display.bind("<Tab>", self._on_scan_tab)
        self.entry_display.bind("<FocusOut>", self._on_focus_out)

        self.verify_btn = tk.Button(
            self, text="VERIFY", bg=ACCENT, fg=WHITE,
            font=("Courier New", 9, "bold"), relief='flat',
            padx=10, pady=5, cursor='hand2',
            command=self._open_verify
        )
        self.verify_btn.grid(row=0, column=1, padx=2)

        self.status_lbl = tk.Label(
            self, text="○ unverified",
            bg=PANEL_BG, fg=RED,
            font=("Courier New", 8)
        )
        self.status_lbl.grid(row=0, column=2, padx=6)

    def _on_focus_out(self, event=None):
        if not self._locked and self.display_var.get().strip():
            self._normalize_current_value()
            self.border_frame.config(bg=AMBER)
            self.status_lbl.config(text="◐ press VERIFY", fg=AMBER)

    def _on_scan_tab(self, event=None):
        if not self._locked and self.display_var.get().strip():
            self._open_verify()
            return "break"
        return None

    def _normalize_current_value(self):
        current_val = self.display_var.get().strip()
        self.last_raw_value = current_val
        normalized = normalize_scanned_value(
            current_val,
            field_label=self.label_text,
            expected_length=self.length,
            any_length=self.any_length
        )
        if normalized and normalized != current_val:
            self.display_var.set(normalized)
        return normalized or current_val

    def _set_prefill(self, value):
        self._prefill = value
        self.last_raw_value = str(value or '')
        self.display_var.set(normalize_scanned_value(
            value,
            field_label=self.label_text,
            expected_length=self.length,
            any_length=self.any_length
        ))
        self.border_frame.config(bg=AMBER)
        self.status_lbl.config(text="◐ pre-filled — verify", fg=AMBER)

    def _open_verify(self):
        if self._locked:
            return
        current_val = self._normalize_current_value()
        if not current_val:
            self.border_frame.config(bg=RED)
            self.entry_display.focus_set()
            return
        parent_win = self.winfo_toplevel()
        dlg = VerifyDialog(
            parent_win, self.label_text,
            prefill=current_val,
            length=self.length,
            any_length=self.any_length
        )
        parent_win.wait_window(dlg)
        if dlg.result:
            self._lock(dlg.result)
        else:
            if self.display_var.get().strip():
                self.border_frame.config(bg=AMBER)

    def _lock(self, value):
        self._value  = value
        self._locked = True
        self.display_var.set(value)
        self.border_frame.config(bg=GREEN)
        self.entry_display.config(
            state='readonly',
            readonlybackground=LOCKED_BG,
            fg=LOCKED_FG
        )
        self.verify_btn.config(state='disabled', bg=PANEL_BG, fg=TEXT_DIM)
        self.status_lbl.config(text="✓ locked", fg=GREEN)
        if self.on_verified:
            self.on_verified(value)
        if self.next_field:
            self.after(50, self._focus_next)

    def _focus_next(self):
        nf = self.next_field
        if nf is None or not nf.winfo_exists():
            return
        if isinstance(nf, VerifiedField):
            if not nf._locked:
                nf.entry_display.focus_set()
        elif hasattr(nf, 'focus_set'):
            nf.focus_set()

    def unlock(self):
        self._locked  = False
        self._value   = ""
        self._prefill = ""
        self.last_raw_value = ""
        self.display_var.set("")
        self.border_frame.config(bg=RED)
        self.entry_display.config(
            state='normal',
            bg=ENTRY_BG, fg=TEXT
        )
        self.verify_btn.config(state='normal', bg=ACCENT, fg=WHITE)
        self.status_lbl.config(text="○ unverified", fg=RED)

    def set_prefill(self, value):
        if self._locked:
            self.unlock()
        self._set_prefill(value)

    @property
    def value(self):
        return self._value

    @property
    def is_locked(self):
        return self._locked


# ─────────────────────────────────────────────
#  HISTORY PANEL WIDGET
# ─────────────────────────────────────────────
class HistoryPanel(tk.Frame):
    def __init__(self, parent, **kwargs):
        super().__init__(parent, bg=PANEL_BG, relief='flat', bd=0, **kwargs)
        tk.Label(self, text="UNIT HISTORY", bg=PANEL_BG, fg=ACCENT,
                 font=("Courier New", 10, "bold")).pack(anchor='w', padx=10, pady=(10, 4))
        self.content = tk.Label(self, text="Enter serial number\nto check history.",
                                bg=PANEL_BG, fg=TEXT_DIM,
                                font=("Courier New", 9), justify='left',
                                wraplength=210)
        self.content.pack(anchor='w', padx=10, pady=4)
        self.warn_lbl = tk.Label(self, text="", bg=PANEL_BG, fg=RED,
                                 font=("Courier New", 9, "bold"),
                                 wraplength=210, justify='left')
        self.warn_lbl.pack(anchor='w', padx=10)

    def show(self, history, serial):
        if not history:
            self.content.config(
                text=f"Serial: {serial}\n\nNo previous record found.\nFirst time submission.",
                fg=GREEN)
            self.warn_lbl.config(text="")
            return

        date_val = history['date']
        yrs = years_since(date_val)
        yr_str = f"{yrs:.1f} years ago" if yrs is not None else "unknown date"

        text = (f"Serial: {serial}\n\n"
                f"Last tested:\n  {date_val}\n  ({yr_str})\n\n"
                f"Lab #: {history['lab_number']}\n"
                f"CHPS:  {history['chps_number']}\n"
                f"Addr:  {history['address_code']}\n"
                f"File:  {history['source_file']}")
        self.content.config(text=text, fg=TEXT)

        if yrs is not None and yrs < RETEST_PROMPT_YEARS:
            self.warn_lbl.config(
                text=f"⚠ WARNING\nOnly {yrs:.1f} yrs since\nlast test!\n(Min: 3 years)",
                fg=RED)
        elif yrs is not None and yrs < RETEST_MIN_YEARS:
            self.warn_lbl.config(
                text=f"✓ No retest prompt\n({yrs:.1f} yrs since last test)",
                fg=GREEN)
        else:
            self.warn_lbl.config(text="✓ Interval OK", fg=GREEN)

    def clear(self):
        self.content.config(text="Enter serial number\nto check history.", fg=TEXT_DIM)
        self.warn_lbl.config(text="")


# ─────────────────────────────────────────────
#  LAST PROCESSED PANEL
# ─────────────────────────────────────────────
class LastProcessedPanel(tk.Frame):
    def __init__(self, parent, wb, **kwargs):
        super().__init__(parent, bg=PANEL_BG, relief='flat', bd=0, **kwargs)
        tk.Label(self, text="LAST PROCESSED", bg=PANEL_BG, fg=ACCENT,
                 font=("Courier New", 10, "bold")).pack(anchor='w', padx=10, pady=(10, 4))
        self._wb = wb
        self.lbl = tk.Label(self, text="", bg=PANEL_BG, fg=TEXT,
                            font=("Courier New", 9), justify='left')
        self.lbl.pack(anchor='w', padx=10, pady=4)
        self.refresh()

    def refresh(self):
        lines = []
        for sheet, prefix in [('RADAR', 'AS26-'), ('LIDAR', 'ASL26-')]:
            try:
                ws = self._wb[sheet]
                for row in range(ws.max_row, 1, -1):
                    v = ws.cell(row=row, column=1).value
                    if v:
                        s = ws.cell(row=row, column=2).value
                        lines.append(f"{sheet}:\n  {v}\n  ({s})\n")
                        break
                else:
                    lines.append(f"{sheet}: none yet\n")
            except Exception:
                lines.append(f"{sheet}: error\n")
        self.lbl.config(text="\n".join(lines))


# ─────────────────────────────────────────────
#  MAIN APPLICATION
# ─────────────────────────────────────────────
class EntryApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("RADAR / LIDAR CERTIFICATE ENTRY SYSTEM")
        self.configure(bg=BG)
        self.resizable(True, True)
        self.minsize(900, 700)

        if not startup_file_check(self):
            self.destroy()
            return

        try:
            self.area_result = load_area_codes()
        except Exception as e:
            messagebox.showerror("Startup Error", f"Cannot load area_code.docx:\n{e}")
            self.destroy(); return

        try:
            self.wb = load_workbook(EXCEL_2026)
        except Exception as e:
            messagebox.showerror("Startup Error", f"Cannot load week_report_2026.xlsx:\n{e}")
            self.destroy(); return

        self.current_date = datetime.now().strftime("%m/%d/%Y")
        self.session_units = load_session()
        self.next_lab_numbers = {'RADAR': None, 'LIDAR': None}
        self._build_ui()
        if self.session_units:
            self.shipping_tab_btn.config(
                text=f"  SHIPPING ({len(self.session_units)})  ")
        self.after(200, self._start_index_build)

    def get_next_lab_number(self, sheet_name, excel_next):
        manual_next = self.next_lab_numbers.get(sheet_name)
        if manual_next is None:
            return excel_next
        return manual_next

    def set_next_lab_number(self, sheet_name, current_number):
        self.next_lab_numbers[sheet_name] = int(current_number) + 1

    def _start_index_build(self):
        import threading
        def _run():
            def _progress(msg):
                try:
                    self.after(0, lambda m=msg: self._update_index_status(m))
                except Exception:
                    pass
            _build_history_index(on_progress=_progress)
        self._update_index_status("Building search index...")
        t = threading.Thread(target=_run, daemon=True)
        t.start()

    def _update_index_status(self, msg):
        try:
            self.search_form.set_index_status(msg)
        except Exception:
            pass

    def _build_ui(self):
        title_bar = tk.Frame(self, bg=ACCENT, height=46)
        title_bar.pack(fill='x')
        title_bar.pack_propagate(False)
        tk.Label(title_bar,
                 text="  RADAR / LIDAR  CERTIFICATE ENTRY SYSTEM",
                 bg=ACCENT, fg=WHITE,
                 font=("Courier New", 13, "bold")).pack(side='left', padx=10)
        tk.Label(title_bar, text=f"  {self.current_date}",
                 bg=ACCENT, fg=WHITE,
                 font=("Courier New", 11)).pack(side='right', padx=16)

        body = tk.Frame(self, bg=BG)
        body.pack(fill='both', expand=True, padx=0, pady=0)

        side = tk.Frame(body, bg=PANEL_BG, width=240)
        side.pack(side='right', fill='y', padx=(0, 0))
        side.pack_propagate(False)

        separator = tk.Frame(side, bg=ACCENT, width=2)
        separator.pack(side='left', fill='y')

        side_inner = tk.Frame(side, bg=PANEL_BG)
        side_inner.pack(side='left', fill='both', expand=True)

        self.last_panel = LastProcessedPanel(side_inner, self.wb)
        self.last_panel.pack(fill='x', padx=4, pady=(0, 6))

        tk.Frame(side_inner, bg=BORDER, height=1).pack(fill='x', padx=10, pady=4)

        self.history_panel = HistoryPanel(side_inner)
        self.history_panel.pack(fill='x', padx=4)

        form_outer = tk.Frame(body, bg=BG)
        form_outer.pack(side='left', fill='both', expand=True)

        tab_bar = tk.Frame(form_outer, bg=BG)
        tab_bar.pack(fill='x', padx=20, pady=(16, 0))

        self.radar_tab_btn = tk.Button(
            tab_bar, text="  RADAR  ", bg=ACCENT, fg=WHITE,
            font=("Courier New", 11, "bold"), relief='flat',
            padx=10, pady=6, cursor='hand2',
            command=lambda: self._switch_tab('RADAR'))
        self.radar_tab_btn.pack(side='left', padx=(0, 4))

        self.lidar_tab_btn = tk.Button(
            tab_bar, text="  LIDAR  ", bg=PANEL_BG, fg=TEXT_DIM,
            font=("Courier New", 11, "bold"), relief='flat',
            padx=10, pady=6, cursor='hand2',
            command=lambda: self._switch_tab('LIDAR'))
        self.lidar_tab_btn.pack(side='left', padx=(0, 4))

        self.shipping_tab_btn = tk.Button(
            tab_bar, text="  SHIPPING  ", bg=PANEL_BG, fg=TEXT_DIM,
            font=("Courier New", 11, "bold"), relief='flat',
            padx=10, pady=6, cursor='hand2',
            command=lambda: self._switch_tab('SHIPPING'))
        self.shipping_tab_btn.pack(side='left', padx=(0, 4))

        self.search_tab_btn = tk.Button(
            tab_bar, text="  SEARCH  ", bg=PANEL_BG, fg=TEXT_DIM,
            font=("Courier New", 11, "bold"), relief='flat',
            padx=10, pady=6, cursor='hand2',
            command=lambda: self._switch_tab('SEARCH'))
        self.search_tab_btn.pack(side='left')

        self.canvas_frame = tk.Frame(form_outer, bg=BG)
        self.canvas_frame.pack(fill='both', expand=True, padx=10, pady=10)

        self.canvas = tk.Canvas(self.canvas_frame, bg=BG, highlightthickness=0)
        scrollbar = ttk.Scrollbar(self.canvas_frame, orient='vertical',
                                  command=self.canvas.yview)
        self.canvas.configure(yscrollcommand=scrollbar.set)
        scrollbar.pack(side='right', fill='y')
        self.canvas.pack(side='left', fill='both', expand=True)

        self.form_frame = tk.Frame(self.canvas, bg=BG)
        self.canvas_window = self.canvas.create_window(
            (0, 0), window=self.form_frame, anchor='nw')

        self.form_frame.bind('<Configure>', self._on_frame_configure)
        self.canvas.bind('<Configure>', self._on_canvas_configure)
        self.canvas.bind_all('<MouseWheel>', self._on_mousewheel)

        self.radar_form = RadarForm(self.form_frame, self)
        self.lidar_form = LidarForm(self.form_frame, self)

        self.shipping_form = ShippingTab(form_outer, self)
        self.search_form   = SearchTab(form_outer, self)

        self._current_tab = 'RADAR'
        self.radar_form.pack(fill='x', padx=10, pady=10)
        self.after(100, self.radar_form.focus_unit_field)

    def _switch_tab(self, tab):
        if self._current_tab == tab:
            return
        self._current_tab = tab
        self.radar_form.pack_forget()
        self.lidar_form.pack_forget()
        self.shipping_form.pack_forget()
        self.search_form.pack_forget()
        for btn in [self.radar_tab_btn, self.lidar_tab_btn,
                    self.shipping_tab_btn, self.search_tab_btn]:
            btn.config(bg=PANEL_BG, fg=TEXT_DIM)

        if tab in ('RADAR', 'LIDAR'):
            self.canvas_frame.pack(fill='both', expand=True, padx=10, pady=10)
        else:
            self.canvas_frame.pack_forget()

        if tab == 'RADAR':
            self.radar_form.pack(fill='x', padx=10, pady=10)
            self.radar_tab_btn.config(bg=ACCENT, fg=WHITE)
            self.radar_form._compute_lab_number()
            self.after(100, self.radar_form.focus_unit_field)
        elif tab == 'LIDAR':
            self.lidar_form.pack(fill='x', padx=10, pady=10)
            self.lidar_tab_btn.config(bg=ACCENT, fg=WHITE)
            self.lidar_form._compute_lab_number()
            self.after(100, self.lidar_form.focus_unit_field)
        elif tab == 'SHIPPING':
            self.shipping_form.refresh()
            self.shipping_form.pack(fill='both', expand=True, padx=10, pady=10)
            self.shipping_tab_btn.config(bg=ACCENT, fg=WHITE)
        elif tab == 'SEARCH':
            self.search_form.pack(fill='both', expand=True, padx=10, pady=10)
            self.search_tab_btn.config(bg=ACCENT, fg=WHITE)
            self.search_form.focus_first()
        self.history_panel.clear()

    def _on_frame_configure(self, event):
        self.canvas.configure(scrollregion=self.canvas.bbox('all'))

    def _on_canvas_configure(self, event):
        self.canvas.itemconfig(self.canvas_window, width=event.width)

    def _on_mousewheel(self, event):
        self.canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

    def show_history(self, serial, sheet_name):
        import threading
        self.history_panel.show(None, serial)
        result_box = [None]
        done_event = threading.Event()

        def _search():
            result_box[0] = search_all_history(serial, sheet_name)
            done_event.set()

        t = threading.Thread(target=_search, daemon=True)
        t.start()

        def _poll():
            if done_event.is_set():
                self.history_panel.show(result_box[0], serial)
                self._history_result = result_box[0]
            else:
                self.after(50, _poll)

        self._history_result = None
        self.after(50, _poll)

        while not done_event.is_set():
            self.update()
        self.history_panel.show(result_box[0], serial)
        return result_box[0]

    def refresh_last_processed(self):
        try:
            self.wb = load_workbook(EXCEL_2026)
            self.last_panel._wb = self.wb
            self.last_panel.refresh()
        except Exception:
            pass

    def add_session_unit(self, unit_info):
        self.session_units.append(unit_info)
        save_session(self.session_units)
        count = len(self.session_units)
        self.shipping_tab_btn.config(
            text=f"  SHIPPING ({count})  ")


# ─────────────────────────────────────────────
#  RADAR FORM
# ─────────────────────────────────────────────
class RadarForm(tk.Frame):
    def __init__(self, parent, app):
        super().__init__(parent, bg=BG)
        self.app = app
        self._history = None
        self._is_retest = False
        self._reused_current_year_lab = False
        self._build()

    def _section(self, text):
        f = tk.Frame(self, bg=BG)
        f.pack(fill='x', pady=(14, 2))
        tk.Label(f, text=f"▸ {text}", bg=BG, fg=ACCENT,
                 font=("Courier New", 10, "bold")).pack(side='left')
        tk.Frame(f, bg=ACCENT, height=1).pack(side='left', fill='x', expand=True, padx=8)

    def _row(self, label, widget):
        f = tk.Frame(self, bg=BG)
        f.pack(fill='x', pady=3)
        tk.Label(f, text=label, bg=BG, fg=TEXT_DIM,
                 font=("Courier New", 9), width=22, anchor='w').pack(side='left')
        widget.pack(side='left')
        return f

    def _build(self):
        # ── AUTO INFO ────────────────────────
        self._section("AUTO-ASSIGNED")
        info_frame = tk.Frame(self, bg=PANEL_BG, relief='flat', bd=0)
        info_frame.pack(fill='x', pady=4, ipady=8)

        self.lab_var  = tk.StringVar(value="—")
        self.date_var = tk.StringVar(value=self.app.current_date)

        r1 = tk.Frame(info_frame, bg=PANEL_BG)
        r1.pack(fill='x', padx=16, pady=2)
        tk.Label(r1, text="Lab Number:", bg=PANEL_BG, fg=TEXT_DIM,
                 font=("Courier New", 9), width=16, anchor='w').pack(side='left')
        self.lab_entry = tk.Entry(r1, textvariable=self.lab_var,
                                  bg=ENTRY_BG, fg=AMBER,
                                  insertbackground=TEXT,
                                  font=("Courier New", 13, "bold"),
                                  width=14, relief='solid', bd=1)
        self.lab_entry.pack(side='left', ipady=3)
        self.lab_entry.bind("<FocusOut>", self._sync_lab_number)
        self.lab_entry.bind("<Return>", self._sync_lab_number)
        tk.Label(r1, text=" editable", bg=PANEL_BG, fg=TEXT_DIM,
                 font=("Courier New", 8)).pack(side='left', padx=6)

        r2 = tk.Frame(info_frame, bg=PANEL_BG)
        r2.pack(fill='x', padx=16, pady=2)
        tk.Label(r2, text="Date:", bg=PANEL_BG, fg=TEXT_DIM,
                 font=("Courier New", 9), width=16, anchor='w').pack(side='left')
        tk.Label(r2, textvariable=self.date_var, bg=PANEL_BG, fg=TEXT,
                 font=("Courier New", 11)).pack(side='left')

        self.retest_lbl = tk.Label(info_frame, text="", bg=PANEL_BG, fg=AMBER,
                                   font=("Courier New", 9, "bold"))
        self.retest_lbl.pack(anchor='w', padx=16)

        # ── UNIT TYPE ────────────────────────
        self._section("UNIT DETAILS")

        type_frame = tk.Frame(self, bg=BG)
        type_frame.pack(fill='x', pady=4)
        tk.Label(type_frame, text="Unit Type:", bg=BG, fg=TEXT_DIM,
                 font=("Courier New", 9), width=22, anchor='w').pack(side='left')

        self.unit_type_var = tk.StringVar(value=_last_radar_type)
        rb_frame = tk.Frame(type_frame, bg=BG)
        rb_frame.pack(side='left')
        for t in RADAR_TYPES:
            rb = tk.Radiobutton(rb_frame, text=t, variable=self.unit_type_var, value=t,
                                bg=BG, fg=TEXT, selectcolor=CARD_BG,
                                activebackground=BG, activeforeground=ACCENT,
                                font=("Courier New", 10, "bold"),
                                indicatoron=True, cursor='hand2',
                                command=self._on_type_change)
            rb.pack(side='left', padx=6)

        self._compute_lab_number()

        # ── SERIAL NUMBER ────────────────────
        serial_frame = tk.Frame(self, bg=BG)
        serial_frame.pack(fill='x', pady=4)
        tk.Label(serial_frame, text="Serial Number:", bg=BG, fg=TEXT_DIM,
                 font=("Courier New", 9), width=22, anchor='w').pack(side='left')

        self.serial_field = VerifiedField(
            serial_frame, "Serial Number", length=6,
            on_verified=self._on_serial_verified)
        self.serial_field.pack(side='left')

        # ── CHPS ─────────────────────────────
        chps_frame = tk.Frame(self, bg=BG)
        chps_frame.pack(fill='x', pady=4)
        tk.Label(chps_frame, text="CHPS Number:", bg=BG, fg=TEXT_DIM,
                 font=("Courier New", 9), width=22, anchor='w').pack(side='left')
        self.chps_field = VerifiedField(chps_frame, "CHPS Number", length=5)
        self.chps_field.pack(side='left')

        # ── ADDRESS ──────────────────────────
        self._section("ADDRESS")

        addr_frame = tk.Frame(self, bg=BG)
        addr_frame.pack(fill='x', pady=4)
        tk.Label(addr_frame, text="Address Code:", bg=BG, fg=TEXT_DIM,
                 font=("Courier New", 9), width=22, anchor='w').pack(side='left')

        self.addr_border = tk.Frame(addr_frame, bg=RED, padx=2, pady=2)
        self.addr_border.pack(side='left')
        self.addr_var = tk.StringVar()
        self.addr_entry = tk.Entry(self.addr_border, textvariable=self.addr_var,
                              bg=ENTRY_BG, fg=TEXT, insertbackground=TEXT,
                              font=("Courier New", 13, "bold"),
                              width=12, relief='flat', bd=4)
        self.addr_entry.pack(ipady=5)
        self.addr_entry.bind('<FocusOut>', self._lookup_address)
        self.addr_entry.bind('<Return>',   self._addr_enter)
        self.addr_entry.bind('<KP_Enter>', self._addr_enter)
        self.addr_entry.bind('<Tab>',      self._addr_enter)

        self.addr_display = tk.Label(self, text="", bg=BG, fg=GREEN,
                                     font=("Courier New", 9), justify='left')
        self.addr_display.pack(anchor='w', padx=24, pady=2)

        # ── FORKS ────────────────────────────
        # FIX 1: Use a proper Button-style checkbox that responds to Enter/Space
        self._section("FORKS (optional)")

        fork_check_frame = tk.Frame(self, bg=BG)
        fork_check_frame.pack(fill='x', pady=4)
        self.fork_var = tk.BooleanVar(value=False)

        # FIX 1: Replaced Checkbutton with a custom toggle button so Enter/Space work reliably
        self.fork_toggle_btn = tk.Button(
            fork_check_frame,
            text="[ ] Tuning fork certificates present  (SPACE/ENTER)",
            bg=BG, fg=TEXT_DIM,
            font=("Courier New", 10), relief='flat',
            padx=0, pady=2, cursor='hand2',
            anchor='w',
            command=self._toggle_forks_btn
        )
        self.fork_toggle_btn.pack(side='left')
        # Bind Enter AND Space so keyboard nav works fully
        self.fork_toggle_btn.bind("<Return>", lambda e: self._toggle_forks_btn())
        self.fork_toggle_btn.bind("<space>", lambda e: self._toggle_forks_btn())

        # Fork fields container
        self.fork_fields_frame = tk.Frame(self, bg=BG)

        fa_row = tk.Frame(self.fork_fields_frame, bg=BG)
        fa_row.pack(fill='x', pady=3)
        tk.Label(fa_row, text="FA Number:", bg=BG, fg=TEXT_DIM,
                 font=("Courier New", 9), width=22, anchor='w').pack(side='left')
        self.fa_field = VerifiedField(fa_row, "FA Number", length=6)
        self.fa_field.pack(side='left')

        fb_row = tk.Frame(self.fork_fields_frame, bg=BG)
        fb_row.pack(fill='x', pady=3)
        tk.Label(fb_row, text="FB Number:", bg=BG, fg=TEXT_DIM,
                 font=("Courier New", 9), width=22, anchor='w').pack(side='left')
        self.fb_field = VerifiedField(fb_row, "FB Number", length=6)
        self.fb_field.pack(side='left')

        # ── ANTENNAS ─────────────────────────
        self.antenna_section = tk.Frame(self, bg=BG)
        self.antenna_section.pack(fill='x')

        ant_hdr = tk.Frame(self.antenna_section, bg=BG)
        ant_hdr.pack(fill='x', pady=(14, 2))
        tk.Label(ant_hdr, text="\u25b8 ANTENNAS", bg=BG, fg=ACCENT,
                 font=("Courier New", 10, "bold")).pack(side='left')
        tk.Frame(ant_hdr, bg=ACCENT, height=1).pack(side='left', fill='x', expand=True, padx=8)

        self.antenna_frame = tk.Frame(self.antenna_section, bg=BG)
        self.antenna_frame.pack(fill='x')

        ant1_row = tk.Frame(self.antenna_frame, bg=BG)
        ant1_row.pack(fill='x', pady=3)
        tk.Label(ant1_row, text="Antenna 1 Serial:", bg=BG, fg=TEXT_DIM,
                 font=("Courier New", 9), width=22, anchor='w').pack(side='left')
        self.ant1_field = VerifiedField(ant1_row, "Antenna 1 Serial", length=6)
        self.ant1_field.pack(side='left')

        ant2_row = tk.Frame(self.antenna_frame, bg=BG)
        ant2_row.pack(fill='x', pady=3)
        tk.Label(ant2_row, text="Antenna 2 Serial:", bg=BG, fg=TEXT_DIM,
                 font=("Courier New", 9), width=22, anchor='w').pack(side='left')
        self.ant2_field = VerifiedField(ant2_row, "Antenna 2 Serial", length=6)
        self.ant2_field.pack(side='left')

        # ── SUBMIT ───────────────────────────
        btn_bar = tk.Frame(self, bg=BG)
        btn_bar.pack(fill='x', pady=20)

        self.submit_btn = tk.Button(
            btn_bar, text="  GENERATE CERTIFICATE & SAVE  ",
            bg=GREEN, fg=TEXT_DARK,
            font=("Courier New", 11, "bold"), relief='flat',
            padx=16, pady=10, cursor='hand2',
            command=self._submit, state='normal')
        self.submit_btn.pack(side='left', padx=4)

        tk.Button(btn_bar, text="CLEAR FORM", bg=PANEL_BG, fg=TEXT_DIM,
                  font=("Courier New", 10), relief='flat',
                  padx=12, pady=10, cursor='hand2',
                  command=self._clear).pack(side='left', padx=4)

        self.status_lbl = tk.Label(self, text="", bg=BG, fg=TEXT_DIM,
                                   font=("Courier New", 9))
        self.status_lbl.pack(pady=4)

        self._after_addr_target = self.ant1_field
        self.after(10, self._wire_tab_order)

    def focus_unit_field(self):
        self.after(50, lambda: self.serial_field.entry_display.focus_set())

    # FIX 1: New toggle method for the button-style fork control
    def _toggle_forks_btn(self):
        """Toggle the fork_var and update button label; called by Enter/Space/Click."""
        new_val = not self.fork_var.get()
        self.fork_var.set(new_val)
        if new_val:
            self.fork_toggle_btn.config(
                text="[✓] Tuning fork certificates present  (SPACE/ENTER)",
                fg=GREEN)
        else:
            self.fork_toggle_btn.config(
                text="[ ] Tuning fork certificates present  (SPACE/ENTER)",
                fg=TEXT_DIM)
        self._toggle_forks()

    def _wire_tab_order(self):
        self.submit_btn.bind("<Return>", lambda e: self._submit())
        self._rewire_tab_order()

    def _rewire_tab_order(self):
        utype   = self.unit_type_var.get()
        is_as   = utype in ("AS", "ASP")
        has_fork = self.fork_var.get()

        self.serial_field.next_field = self.chps_field
        self.chps_field.next_field   = self.addr_entry

        if has_fork:
            self.fa_field.next_field = self.fb_field
            if not is_as:
                self.fb_field.next_field = self.ant1_field
                self.ant1_field.next_field = self.ant2_field
                self.ant2_field.next_field = self.submit_btn
                self._after_addr_target = self.fa_field
            else:
                self.fb_field.next_field = self.submit_btn
                self._after_addr_target = self.fa_field
        else:
            if not is_as:
                self.ant1_field.next_field = self.ant2_field
                self.ant2_field.next_field = self.submit_btn
                self._after_addr_target = self.ant1_field
            else:
                self._after_addr_target = self.submit_btn

    # ── LOGIC ────────────────────────────────

    def _compute_lab_number(self):
        try:
            wb = load_workbook(EXCEL_2026, read_only=True, data_only=True)
            ws = wb['RADAR']
            excel_next  = get_last_lab_number(ws, RADAR_PREFIX) + 1
            wb.close()
        except Exception:
            ws = self.app.wb['RADAR']
            excel_next  = get_last_lab_number(ws, RADAR_PREFIX) + 1
        n = self.app.get_next_lab_number('RADAR', excel_next)
        self._lab_number = str(n)
        self.lab_var.set(format_lab_number(RADAR_PREFIX, n))

    def _sync_lab_number(self, event=None):
        n = parse_lab_number(self.lab_var.get(), RADAR_PREFIX)
        if n is None:
            self.lab_entry.config(fg=RED)
            return False
        self._lab_number = str(n)
        self.lab_var.set(format_lab_number(RADAR_PREFIX, n))
        self.lab_entry.config(fg=AMBER)
        return True

    def _reuse_current_year_failed_lab(self, history):
        if not is_current_year_failed_record(history):
            return False
        n = parse_lab_number(history.get('lab_number'), RADAR_PREFIX)
        if n is None:
            return False
        self._lab_number = str(n)
        self.lab_var.set(format_lab_number(RADAR_PREFIX, n))
        self._reused_current_year_lab = True
        return True

    def _on_type_change(self, event=None):
        global _last_radar_type
        utype = self.unit_type_var.get()
        _last_radar_type = utype
        is_as = utype in ("AS", "ASP")
        if is_as:
            self.antenna_section.pack_forget()
        else:
            self.antenna_section.pack(fill='x')
        self._rewire_tab_order()

    def _on_serial_verified(self, serial):
        detected_type = detect_unit_type_prefix(self.serial_field.last_raw_value)
        if detected_type in RADAR_TYPES and detected_type != self.unit_type_var.get():
            self.unit_type_var.set(detected_type)
            self._on_type_change()
        elif detected_type in LIDAR_TYPES:
            messagebox.showwarning(
                "LIDAR Scan Detected",
                f"Scanned value looks like LIDAR type {detected_type}.\n"
                f"Use the LIDAR tab for this unit.")
        history = self.app.show_history(serial, 'RADAR')
        self._history = history
        self._is_retest = False
        self._reused_current_year_lab = False
        if history:
            if self._reuse_current_year_failed_lab(history):
                self._apply_retest_prefill(history, mark_retest=True)
                self.retest_lbl.config(
                    text="◐ RETEST — same-year failed unit; fields pre-filled")
                return
            yrs = years_since(history['date'])

            if yrs is not None and yrs < RETEST_PROMPT_YEARS:
                ans = messagebox.askyesno(
                    "⚠ Early Retest Warning",
                    f"This unit was last tested {yrs:.1f} years ago\n"
                    f"(Lab: {history['lab_number']}, Date: {history['date']})\n\n"
                    f"Minimum interval is {RETEST_MIN_YEARS:.0f} years.\n\n"
                    f"Is this a RETEST (proceed with auto-fill)?")
                if ans:
                    self._apply_retest_prefill(history, mark_retest=True)
                else:
                    self._history = None
                    self.retest_lbl.config(text="")
                    return
            elif yrs is not None and yrs < RETEST_MIN_YEARS:
                self.retest_lbl.config(
                    text=f"✓ Previous test {yrs:.1f} yrs ago — auto-filled")
                self._apply_retest_prefill(history, mark_retest=False)
            else:
                self._apply_retest_prefill(history, mark_retest=False)
        else:
            self.retest_lbl.config(text="")

    def _apply_retest_prefill(self, history, mark_retest=True):
        """Populate all available fields from the history record."""
        self._is_retest = mark_retest
        if mark_retest:
            self.retest_lbl.config(text="◐ RETEST — fields pre-filled (verify to confirm)")
        else:
            self.retest_lbl.config(text="◐ Previous entry — fields auto-filled; enter current CHPS")
        if mark_retest and history.get('chps_number'):
            self.chps_field.set_prefill(history['chps_number'])
        if history.get('address_code'):
            self.addr_var.set(history['address_code'])
            self._lookup_address()
        # FIX 3: Always fill antenna fields if available
        if history.get('antenna1_number'):
            self.ant1_field.set_prefill(history['antenna1_number'])
        if history.get('antenna2_number'):
            self.ant2_field.set_prefill(history['antenna2_number'])

    def _addr_enter(self, event=None):
        self._lookup_address()
        if self._address_1:
            target = getattr(self, '_after_addr_target', self.submit_btn)
            if isinstance(target, VerifiedField):
                self.after(50, lambda: target.entry_display.focus_set())
            elif target == self.fork_toggle_btn or hasattr(target, 'focus_set'):
                self.after(50, target.focus_set)
        if event and getattr(event, 'keysym', '') == 'Tab':
            return "break"

    def _lookup_address(self, event=None):
        code = normalize_address_code(self.addr_var.get())
        if code != self.addr_var.get().strip():
            self.addr_var.set(code)
        if not code:
            return
        key = f'CHP ({code})'
        if key in self.app.area_result:
            ind = self.app.area_result.index(key)
            a1  = self.app.area_result[ind + 2]
            a2  = self.app.area_result[ind + 3]
            self.addr_display.config(text=f"  {a1}\n  {a2}", fg=GREEN)
            self.addr_border.config(bg=GREEN)
            self._address_1 = a1
            self._address_2 = a2
        else:
            self.addr_display.config(text=f"  ⚠ Code '{code}' not found.", fg=RED)
            self.addr_border.config(bg=RED)
            self._address_1 = None
            self._address_2 = None

    def _toggle_forks(self):
        """Show/hide the fork fields frame based on fork_var state."""
        if self.fork_var.get():
            self.fork_fields_frame.pack(fill='x', pady=2,
                                        after=self.fork_toggle_btn.master)
            if not self.fa_field.value and not self.fa_field.display_var.get().strip():
                self.fa_field.set_prefill(DEFAULT_FA)
            if not self.fb_field.value and not self.fb_field.display_var.get().strip():
                self.fb_field.set_prefill(DEFAULT_FB)
            self._rewire_tab_order()
            # Move focus to FA field after toggle
            self.after(60, lambda: self.fa_field.entry_display.focus_set())
        else:
            self.fork_fields_frame.pack_forget()
            self._rewire_tab_order()

    def _validate(self):
        utype = self.unit_type_var.get()
        is_as = utype in ("AS", "ASP")
        errors = []
        if not self._sync_lab_number():
            errors.append("Lab Number is invalid")
        if not self.serial_field.is_locked:
            errors.append("Serial Number not verified")
        if not self.chps_field.is_locked:
            errors.append("CHPS Number not verified")
        if not self.addr_var.get().strip():
            errors.append("Address Code is empty")
        if not hasattr(self, '_address_1') or not self._address_1:
            errors.append("Address Code not found / not looked up")
        if self.fork_var.get():
            if not self.fa_field.is_locked:
                errors.append("FA Number not verified")
            if not self.fb_field.is_locked:
                errors.append("FB Number not verified")
        if not is_as:
            if not self.ant1_field.is_locked:
                errors.append("Antenna 1 not verified")
            if not self.ant2_field.is_locked:
                errors.append("Antenna 2 not verified")
        return errors

    def _submit(self):
        errors = self._validate()
        if errors:
            messagebox.showerror("Cannot Submit",
                                 "Please fix the following:\n\n• " +
                                 "\n• ".join(errors))
            return

        utype  = self.unit_type_var.get()
        is_as  = utype in ("AS", "ASP")
        serial = self.serial_field.value
        chps   = self.chps_field.value
        addr   = self.addr_var.get().strip()
        is_retest = self._is_retest

        fa_number = self.fa_field.value if self.fork_var.get() else DEFAULT_FA
        fb_number = self.fb_field.value if self.fork_var.get() else DEFAULT_FB
        ant1 = self.ant1_field.value if not is_as else "N/A"
        ant2 = self.ant2_field.value if not is_as else "N/A"

        tpl_name = RADAR_TEMPLATES[utype]
        files_needed = [
            ("Excel Log (week_report_2026.xlsx)", EXCEL_2026),
            (f"RADAR Template ({tpl_name})",
             os.path.join(TEMPLATE_DIR, tpl_name)),
        ]
        if not wait_for_files_free(self.winfo_toplevel(), files_needed):
            self.status_lbl.config(
                text="⚠  Save cancelled — values preserved.", fg=AMBER)
            return

        try:
            tpl_path = os.path.join(TEMPLATE_DIR, RADAR_TEMPLATES[utype])
            doc = DocxTemplate(tpl_path)
            context = {
                'date': self.app.current_date,
                'lab_number': self._lab_number,
                'chps_number': chps,
                'serial_number': serial,
                'fa_number': fa_number,
                'fb_number': fb_number,
                'antenna1_number': ant1,
                'antenna2_number': ant2,
                'address_code': addr,
                'st_address': self._address_1,
                'area_address': self._address_2,
            }
            doc.render(context)
            out_path = os.path.join(RADAR_OUT_DIR, f"AS26-{self._lab_number}.docx")
            doc.save(out_path)

            ws = self.app.wb['RADAR']
            lab_display = f'AS26-{self._lab_number} (RETEST)' if is_retest else f'AS26-{self._lab_number}'
            utype_prefix = utype[:2]
            unit_name = RADAR_UNIT_NAMES[utype]

            if not is_as:
                ws.append([lab_display, utype_prefix+serial, 'CHPS'+chps, addr,
                           self.app.current_date, " ", " ", " ", unit_name,
                           " ", " ", "S/N "+ant1, "S/N "+ant2])
            else:
                ws.append([lab_display, utype_prefix+serial, 'CHPS'+chps, addr,
                           self.app.current_date, " ", " ", " ", unit_name,
                           " ", " ", "N/A", "N/A"])

            self.app.wb.save(EXCEL_2026)
            update_history_index_record('RADAR', {
                'lab_number':      lab_display,
                'unit_serial':     utype_prefix + serial,
                'chps_number':     chps,
                'address_code':    normalize_address_code(addr),
                'date':            self.app.current_date,
                'shipped_date':    " ",
                'antenna1_number': ant1 if not is_as else None,
                'antenna2_number': ant2 if not is_as else None,
                'source_file':     os.path.basename(EXCEL_2026),
                'source_path':     EXCEL_2026,
            })

            self.status_lbl.config(
                text=f"✓ Certificate saved: AS26-{self._lab_number}",
                fg=GREEN)
            messagebox.showinfo("Success",
                                f"Certificate generated:\n  AS26-{self._lab_number}\n\n"
                                f"Saved to Excel.")
            self.app.add_session_unit({
                'type': 'RADAR',
                'lab_number': f'AS26-{self._lab_number}',
                'address_code': addr,
                'address_1': self._address_1,
                'address_2': self._address_2,
                'unit_name': RADAR_UNIT_NAMES[utype],
                'serial': serial,
                'chps': chps,
            })
            if not self._reused_current_year_lab:
                self.app.set_next_lab_number('RADAR', int(self._lab_number))
            self.app.refresh_last_processed()
            self._clear()

        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate certificate:\n{e}")

    def _clear(self):
        self._history = None
        self._is_retest = False
        self._reused_current_year_lab = False
        self._address_1 = None
        self._address_2 = None
        self.unit_type_var.set(_last_radar_type)
        self.serial_field.unlock()
        self.chps_field.unlock()
        self.addr_var.set("")
        self.addr_display.config(text="")
        self.fa_field.unlock()
        self.fb_field.unlock()
        self.ant1_field.unlock()
        self.ant2_field.unlock()
        self.fork_var.set(False)
        # FIX 1: Reset the toggle button text on clear
        self.fork_toggle_btn.config(
            text="[ ] Tuning fork certificates present  (SPACE/ENTER)",
            fg=TEXT_DIM)
        self.fork_fields_frame.pack_forget()
        self.antenna_frame.pack(fill='x')
        self.retest_lbl.config(text="")
        self.status_lbl.config(text="")
        self._compute_lab_number()
        self.app.history_panel.clear()
        self.focus_unit_field()


# ─────────────────────────────────────────────
#  LIDAR FORM
# ─────────────────────────────────────────────
class LidarForm(tk.Frame):
    def __init__(self, parent, app):
        super().__init__(parent, bg=BG)
        self.app = app
        self._history = None
        self._is_retest = False
        self._reused_current_year_lab = False
        self._build()

    def _section(self, text):
        f = tk.Frame(self, bg=BG)
        f.pack(fill='x', pady=(14, 2))
        tk.Label(f, text=f"▸ {text}", bg=BG, fg=ACCENT,
                 font=("Courier New", 10, "bold")).pack(side='left')
        tk.Frame(f, bg=ACCENT, height=1).pack(side='left', fill='x', expand=True, padx=8)

    def _build(self):
        # ── AUTO INFO ────────────────────────
        self._section("AUTO-ASSIGNED")
        info_frame = tk.Frame(self, bg=PANEL_BG, relief='flat', bd=0)
        info_frame.pack(fill='x', pady=4, ipady=8)

        self.lab_var  = tk.StringVar(value="—")
        self.date_var = tk.StringVar(value=self.app.current_date)

        r1 = tk.Frame(info_frame, bg=PANEL_BG)
        r1.pack(fill='x', padx=16, pady=2)
        tk.Label(r1, text="Lab Number:", bg=PANEL_BG, fg=TEXT_DIM,
                 font=("Courier New", 9), width=16, anchor='w').pack(side='left')
        self.lab_entry = tk.Entry(r1, textvariable=self.lab_var,
                                  bg=ENTRY_BG, fg=AMBER,
                                  insertbackground=TEXT,
                                  font=("Courier New", 13, "bold"),
                                  width=14, relief='solid', bd=1)
        self.lab_entry.pack(side='left', ipady=3)
        self.lab_entry.bind("<FocusOut>", self._sync_lab_number)
        self.lab_entry.bind("<Return>", self._sync_lab_number)
        tk.Label(r1, text=" editable", bg=PANEL_BG, fg=TEXT_DIM,
                 font=("Courier New", 8)).pack(side='left', padx=6)

        r2 = tk.Frame(info_frame, bg=PANEL_BG)
        r2.pack(fill='x', padx=16, pady=2)
        tk.Label(r2, text="Date:", bg=PANEL_BG, fg=TEXT_DIM,
                 font=("Courier New", 9), width=16, anchor='w').pack(side='left')
        tk.Label(r2, textvariable=self.date_var, bg=PANEL_BG, fg=TEXT,
                 font=("Courier New", 11)).pack(side='left')

        self.retest_lbl = tk.Label(info_frame, text="", bg=PANEL_BG, fg=AMBER,
                                   font=("Courier New", 9, "bold"))
        self.retest_lbl.pack(anchor='w', padx=16)

        # ── UNIT TYPE ────────────────────────
        self._section("UNIT DETAILS")

        type_frame = tk.Frame(self, bg=BG)
        type_frame.pack(fill='x', pady=4)
        tk.Label(type_frame, text="Unit Type:", bg=BG, fg=TEXT_DIM,
                 font=("Courier New", 9), width=22, anchor='w').pack(side='left')

        self.unit_type_var = tk.StringVar(value=_last_lidar_type)
        rb_frame = tk.Frame(type_frame, bg=BG)
        rb_frame.pack(side='left')
        for t in LIDAR_TYPES:
            rb = tk.Radiobutton(rb_frame, text=t, variable=self.unit_type_var, value=t,
                                bg=BG, fg=TEXT, selectcolor=CARD_BG,
                                activebackground=BG, activeforeground=ACCENT,
                                font=("Courier New", 10, "bold"),
                                indicatoron=True, cursor='hand2',
                                command=self._on_lidar_type_change)
            rb.pack(side='left', padx=6)

        self._compute_lab_number()

        # ── SERIAL ───────────────────────────
        serial_frame = tk.Frame(self, bg=BG)
        serial_frame.pack(fill='x', pady=4)
        tk.Label(serial_frame, text="Serial Number:", bg=BG, fg=TEXT_DIM,
                 font=("Courier New", 9), width=22, anchor='w').pack(side='left')
        self.serial_field = VerifiedField(
            serial_frame, "Serial Number", length=6, any_length=False,
            on_verified=self._on_serial_verified)
        self.serial_field.pack(side='left')

        # ── CHPS ─────────────────────────────
        chps_frame = tk.Frame(self, bg=BG)
        chps_frame.pack(fill='x', pady=4)
        tk.Label(chps_frame, text="CHPS Number:", bg=BG, fg=TEXT_DIM,
                 font=("Courier New", 9), width=22, anchor='w').pack(side='left')
        self.chps_field = VerifiedField(chps_frame, "CHPS Number", length=5)
        self.chps_field.pack(side='left')

        # ── ADDRESS ──────────────────────────
        self._section("ADDRESS")

        addr_frame = tk.Frame(self, bg=BG)
        addr_frame.pack(fill='x', pady=4)
        tk.Label(addr_frame, text="Address Code:", bg=BG, fg=TEXT_DIM,
                 font=("Courier New", 9), width=22, anchor='w').pack(side='left')

        self.addr_border = tk.Frame(addr_frame, bg=RED, padx=2, pady=2)
        self.addr_border.pack(side='left')
        self.addr_var = tk.StringVar()
        self.addr_entry = tk.Entry(self.addr_border, textvariable=self.addr_var,
                              bg=ENTRY_BG, fg=TEXT, insertbackground=TEXT,
                              font=("Courier New", 13, "bold"),
                              width=12, relief='flat', bd=4)
        self.addr_entry.pack(ipady=5)
        self.addr_entry.bind('<FocusOut>', self._lookup_address)
        self.addr_entry.bind('<Return>',   self._addr_enter)
        self.addr_entry.bind('<KP_Enter>', self._addr_enter)
        self.addr_entry.bind('<Tab>',      self._addr_enter)

        self.addr_display = tk.Label(self, text="", bg=BG, fg=GREEN,
                                     font=("Courier New", 9), justify='left')
        self.addr_display.pack(anchor='w', padx=24, pady=2)

        # ── SUBMIT ───────────────────────────
        btn_bar = tk.Frame(self, bg=BG)
        btn_bar.pack(fill='x', pady=20)

        self.submit_btn = tk.Button(
            btn_bar, text="  GENERATE CERTIFICATE & SAVE  ",
            bg=GREEN, fg=TEXT_DARK,
            font=("Courier New", 11, "bold"), relief='flat',
            padx=16, pady=10, cursor='hand2',
            command=self._submit)
        self.submit_btn.pack(side='left', padx=4)

        tk.Button(btn_bar, text="CLEAR FORM", bg=PANEL_BG, fg=TEXT_DIM,
                  font=("Courier New", 10), relief='flat',
                  padx=12, pady=10, cursor='hand2',
                  command=self._clear).pack(side='left', padx=4)

        self.status_lbl = tk.Label(self, text="", bg=BG, fg=TEXT_DIM,
                                   font=("Courier New", 9))
        self.status_lbl.pack(pady=4)

        self.after(10, self._wire_tab_order)

    def focus_unit_field(self):
        self.after(50, lambda: self.serial_field.entry_display.focus_set())

    def _wire_tab_order(self):
        self.serial_field.next_field = self.chps_field
        self.chps_field.next_field   = self.addr_entry
        self.submit_btn.bind("<Return>", lambda e: self._submit())

    def _compute_lab_number(self):
        try:
            wb = load_workbook(EXCEL_2026, read_only=True, data_only=True)
            ws = wb['LIDAR']
            excel_next  = get_last_lab_number(ws, LIDAR_PREFIX) + 1
            wb.close()
        except Exception:
            ws = self.app.wb['LIDAR']
            excel_next  = get_last_lab_number(ws, LIDAR_PREFIX) + 1
        n = self.app.get_next_lab_number('LIDAR', excel_next)
        self._lab_number = str(n)
        self.lab_var.set(format_lab_number(LIDAR_PREFIX, n))

    def _sync_lab_number(self, event=None):
        n = parse_lab_number(self.lab_var.get(), LIDAR_PREFIX)
        if n is None:
            self.lab_entry.config(fg=RED)
            return False
        self._lab_number = str(n)
        self.lab_var.set(format_lab_number(LIDAR_PREFIX, n))
        self.lab_entry.config(fg=AMBER)
        return True

    def _reuse_current_year_failed_lab(self, history):
        if not is_current_year_failed_record(history):
            return False
        n = parse_lab_number(history.get('lab_number'), LIDAR_PREFIX)
        if n is None:
            return False
        self._lab_number = str(n)
        self.lab_var.set(format_lab_number(LIDAR_PREFIX, n))
        self._reused_current_year_lab = True
        return True

    def _on_lidar_type_change(self, event=None):
        global _last_lidar_type
        utype = self.unit_type_var.get()
        _last_lidar_type = utype
        if hasattr(self, 'serial_field'):
            self.serial_field.any_length = (utype == 'LP')

    def _on_serial_verified(self, serial):
        detected_type = detect_unit_type_prefix(self.serial_field.last_raw_value)
        if detected_type in LIDAR_TYPES and detected_type != self.unit_type_var.get():
            self.unit_type_var.set(detected_type)
            self._on_lidar_type_change()
        elif detected_type in RADAR_TYPES:
            messagebox.showwarning(
                "RADAR Scan Detected",
                f"Scanned value looks like RADAR type {detected_type}.\n"
                f"Use the RADAR tab for this unit.")
        history = self.app.show_history(serial, 'LIDAR')
        self._history = history
        self._is_retest = False
        self._reused_current_year_lab = False
        if history:
            if self._reuse_current_year_failed_lab(history):
                self._apply_retest_prefill(history, mark_retest=True)
                self.retest_lbl.config(
                    text="◐ RETEST — same-year failed unit; fields pre-filled")
                return
            yrs = years_since(history['date'])

            if yrs is not None and yrs < RETEST_PROMPT_YEARS:
                ans = messagebox.askyesno(
                    "⚠ Early Retest Warning",
                    f"This unit was last tested {yrs:.1f} years ago\n"
                    f"(Lab: {history['lab_number']}, Date: {history['date']})\n\n"
                    f"Minimum interval is {RETEST_MIN_YEARS:.0f} years.\n\n"
                    f"Is this a RETEST (proceed with auto-fill)?")
                if ans:
                    self._apply_retest_prefill(history, mark_retest=True)
                else:
                    self._history = None
                    self.retest_lbl.config(text="")
                    return
            elif yrs is not None and yrs < RETEST_MIN_YEARS:
                self.retest_lbl.config(
                    text=f"✓ Previous test {yrs:.1f} yrs ago — auto-filled")
                self._apply_retest_prefill(history, mark_retest=False)
            else:
                self._apply_retest_prefill(history, mark_retest=False)
        else:
            self.retest_lbl.config(text="")

    def _apply_retest_prefill(self, history, mark_retest=True):
        """Populate all available fields from the history record."""
        self._is_retest = mark_retest
        if mark_retest:
            self.retest_lbl.config(text="◐ RETEST — fields pre-filled (verify to confirm)")
        else:
            self.retest_lbl.config(text="◐ Previous entry — fields auto-filled; enter current CHPS")
        if mark_retest and history.get('chps_number'):
            self.chps_field.set_prefill(history['chps_number'])
        if history.get('address_code'):
            self.addr_var.set(history['address_code'])
            self._lookup_address()

    def _addr_enter(self, event=None):
        self._lookup_address()
        if self._address_1:
            self.after(50, lambda: self.submit_btn.focus_set())
        if event and getattr(event, 'keysym', '') == 'Tab':
            return "break"

    def _lookup_address(self, event=None):
        code = normalize_address_code(self.addr_var.get())
        if code != self.addr_var.get().strip():
            self.addr_var.set(code)
        if not code:
            return
        key = f'CHP ({code})'
        if key in self.app.area_result:
            ind = self.app.area_result.index(key)
            a1  = self.app.area_result[ind + 2]
            a2  = self.app.area_result[ind + 3]
            self.addr_display.config(text=f"  {a1}\n  {a2}", fg=GREEN)
            self.addr_border.config(bg=GREEN)
            self._address_1 = a1
            self._address_2 = a2
        else:
            self.addr_display.config(text=f"  ⚠ Code '{code}' not found.", fg=RED)
            self.addr_border.config(bg=RED)
            self._address_1 = None
            self._address_2 = None

    def _validate(self):
        errors = []
        if not self._sync_lab_number():
            errors.append("Lab Number is invalid")
        if not self.serial_field.is_locked:
            errors.append("Serial Number not verified")
        if not self.chps_field.is_locked:
            errors.append("CHPS Number not verified")
        if not self.addr_var.get().strip():
            errors.append("Address Code is empty")
        if not hasattr(self, '_address_1') or not self._address_1:
            errors.append("Address Code not found / not looked up")
        return errors

    def _submit(self):
        errors = self._validate()
        if errors:
            messagebox.showerror("Cannot Submit",
                                 "Please fix the following:\n\n• " +
                                 "\n• ".join(errors))
            return

        utype  = self.unit_type_var.get()
        serial = self.serial_field.value
        chps   = self.chps_field.value
        addr   = self.addr_var.get().strip()
        is_retest = self._is_retest

        tpl_name = LIDAR_TEMPLATES[utype]
        files_needed = [
            ("Excel Log (week_report_2026.xlsx)", EXCEL_2026),
            (f"LIDAR Template ({tpl_name})",
             os.path.join(TEMPLATE_DIR, tpl_name)),
        ]
        if not wait_for_files_free(self.winfo_toplevel(), files_needed):
            self.status_lbl.config(
                text="⚠  Save cancelled — values preserved.", fg=AMBER)
            return

        try:
            tpl_path = os.path.join(TEMPLATE_DIR, LIDAR_TEMPLATES[utype])
            doc = DocxTemplate(tpl_path)
            context = {
                'date': self.app.current_date,
                'lab_number': self._lab_number,
                'chps_number': chps,
                'serial_number': serial,
                'address_code': addr,
                'st_address': self._address_1,
                'area_address': self._address_2,
            }
            doc.render(context)
            out_path = os.path.join(LIDAR_OUT_DIR, f"ASL26-{self._lab_number}.docx")
            doc.save(out_path)

            ws = self.app.wb['LIDAR']
            lab_display = (f'ASL26-{self._lab_number} (RETEST)'
                           if is_retest else f'ASL26-{self._lab_number}')
            unit_name = LIDAR_UNIT_NAMES[utype]
            ws.append([lab_display, utype+serial, 'CHPS'+chps, addr,
                       self.app.current_date, " ", " ", " ", unit_name])
            self.app.wb.save(EXCEL_2026)
            update_history_index_record('LIDAR', {
                'lab_number':      lab_display,
                'unit_serial':     utype + serial,
                'chps_number':     chps,
                'address_code':    normalize_address_code(addr),
                'date':            self.app.current_date,
                'shipped_date':    " ",
                'antenna1_number': None,
                'antenna2_number': None,
                'source_file':     os.path.basename(EXCEL_2026),
                'source_path':     EXCEL_2026,
            })

            self.status_lbl.config(
                text=f"✓ Certificate saved: ASL26-{self._lab_number}", fg=GREEN)
            messagebox.showinfo("Success",
                                f"Certificate generated:\n  ASL26-{self._lab_number}\n\n"
                                f"Saved to Excel.")
            self.app.add_session_unit({
                'type': 'LIDAR',
                'lab_number': f'ASL26-{self._lab_number}',
                'address_code': addr,
                'address_1': self._address_1,
                'address_2': self._address_2,
                'unit_name': LIDAR_UNIT_NAMES[utype],
                'serial': serial,
                'chps': chps,
            })
            if not self._reused_current_year_lab:
                self.app.set_next_lab_number('LIDAR', int(self._lab_number))
            self.app.refresh_last_processed()
            self._clear()

        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate certificate:\n{e}")

    def _clear(self):
        self._history = None
        self._is_retest = False
        self._reused_current_year_lab = False
        self._address_1 = None
        self._address_2 = None
        self.unit_type_var.set(_last_lidar_type)
        self.serial_field.unlock()
        self.chps_field.unlock()
        self.addr_var.set("")
        self.addr_display.config(text="")
        self.retest_lbl.config(text="")
        self.status_lbl.config(text="")
        self._compute_lab_number()
        self.app.history_panel.clear()
        self.focus_unit_field()


# ─────────────────────────────────────────────
#  MANUAL ADD DIALOG
# ─────────────────────────────────────────────
class ManualAddDialog(tk.Toplevel):
    def __init__(self, parent, app, on_complete):
        super().__init__(parent)
        self.app         = app
        self.on_complete = on_complete
        self.title("Manual Add — Shipping")
        self.configure(bg=WHITE)
        self.resizable(False, False)
        self.grab_set()
        self.focus_set()

        self._added_units = []
        self._current_idx = 0
        self._total       = 0
        self._unit_type   = 'RADAR'

        self._center()
        self._build_step1()

    def _center(self):
        self.update_idletasks()
        pw = self.master.winfo_rootx()
        ph = self.master.winfo_rooty()
        px = pw + (self.master.winfo_width()  - 480) // 2
        py = ph + (self.master.winfo_height() - 400) // 2
        self.geometry(f"480x420+{px}+{py}")

    def _clear(self):
        for w in self.winfo_children():
            w.destroy()

    def _build_step1(self):
        self._clear()
        self.configure(bg=WHITE)

        tk.Label(self, text="MANUAL ADD TO SHIPPING",
                 bg=WHITE, fg=ACCENT,
                 font=("Courier New", 12, "bold")).pack(pady=(20, 4))
        tk.Label(self, text="Add units that were processed outside this session.",
                 bg=WHITE, fg=TEXT_DIM,
                 font=("Courier New", 9)).pack(pady=(0, 16))

        type_frame = tk.Frame(self, bg=WHITE)
        type_frame.pack(pady=6)
        tk.Label(type_frame, text="Unit Type:", bg=WHITE, fg=TEXT,
                 font=("Courier New", 10), width=14, anchor='w').pack(side='left')
        self._type_var = tk.StringVar(value='RADAR')
        for val, lbl in [('RADAR', 'RADAR'), ('LIDAR', 'LIDAR')]:
            tk.Radiobutton(type_frame, text=lbl, variable=self._type_var,
                           value=val, bg=WHITE, fg=TEXT, selectcolor=PANEL_BG,
                           font=("Courier New", 10),
                           activebackground=WHITE).pack(side='left', padx=10)

        num_frame = tk.Frame(self, bg=WHITE)
        num_frame.pack(pady=6)
        tk.Label(num_frame, text="Number of units:", bg=WHITE, fg=TEXT,
                 font=("Courier New", 10), width=14, anchor='w').pack(side='left')
        self._num_var = tk.StringVar(value='1')
        num_entry = tk.Entry(num_frame, textvariable=self._num_var,
                             bg=PANEL_BG, fg=TEXT, insertbackground=TEXT,
                             font=("Courier New", 13, "bold"),
                             width=6, relief='solid', bd=1)
        num_entry.pack(side='left', ipady=4)
        num_entry.focus_set()

        self._msg = tk.Label(self, text="", bg=WHITE, fg=RED,
                             font=("Courier New", 9))
        self._msg.pack(pady=4)

        btn_f = tk.Frame(self, bg=WHITE)
        btn_f.pack(pady=16)
        tk.Button(btn_f, text="NEXT →", bg=ACCENT, fg=WHITE,
                  font=("Courier New", 10, "bold"), relief='flat',
                  padx=16, pady=7, cursor='hand2',
                  command=self._step1_next).pack(side='left', padx=6)
        tk.Button(btn_f, text="CANCEL", bg=PANEL_BG, fg=TEXT,
                  font=("Courier New", 10), relief='flat',
                  padx=16, pady=7, cursor='hand2',
                  command=self.destroy).pack(side='left', padx=6)

        num_entry.bind("<Return>", lambda e: self._step1_next())

    def _step1_next(self):
        try:
            n = int(self._num_var.get().strip())
            if n < 1:
                raise ValueError
        except ValueError:
            self._msg.config(text="⚠  Enter a valid number of units (1 or more).")
            return
        self._total     = n
        self._unit_type = self._type_var.get()
        self._current_idx = 0
        self._added_units = []
        self._build_step2()

    def _build_step2(self):
        self._clear()
        idx   = self._current_idx
        total = self._total
        utype = self._unit_type
        prefix = 'AS26-' if utype == 'RADAR' else 'ASL26-'

        tk.Label(self,
                 text=f"UNIT {idx + 1} OF {total}  —  {utype}",
                 bg=WHITE, fg=ACCENT,
                 font=("Courier New", 12, "bold")).pack(pady=(20, 4))

        lab_frame = tk.Frame(self, bg=WHITE)
        lab_frame.pack(pady=8)
        tk.Label(lab_frame,
                 text=f"Lab / Log Number (e.g. {prefix}105):",
                 bg=WHITE, fg=TEXT,
                 font=("Courier New", 10), width=18, anchor='w',
                 justify='left').pack(side='left')
        self._lab_var = tk.StringVar()
        lab_entry = tk.Entry(lab_frame, textvariable=self._lab_var,
                             bg=PANEL_BG, fg=TEXT, insertbackground=TEXT,
                             font=("Courier New", 13, "bold"),
                             width=14, relief='solid', bd=1)
        lab_entry.pack(side='left', ipady=5)
        lab_entry.focus_set()
        lab_entry.bind("<Return>", lambda e: self._addr_entry.focus_set())

        addr_frame = tk.Frame(self, bg=WHITE)
        addr_frame.pack(pady=8)
        tk.Label(addr_frame, text="Address Code:",
                 bg=WHITE, fg=TEXT,
                 font=("Courier New", 10), width=18, anchor='w').pack(side='left')
        self._addr_var = tk.StringVar()
        self._addr_entry = tk.Entry(addr_frame, textvariable=self._addr_var,
                                    bg=PANEL_BG, fg=TEXT, insertbackground=TEXT,
                                    font=("Courier New", 13, "bold"),
                                    width=14, relief='solid', bd=1)
        self._addr_entry.pack(side='left', ipady=5)
        self._addr_entry.bind("<Return>", lambda e: self._check_and_preview())

        self._addr_display = tk.Label(self, text="",
                                      bg=WHITE, fg=GREEN,
                                      font=("Courier New", 9))
        self._addr_display.pack()

        self._msg2 = tk.Label(self, text="", bg=WHITE, fg=RED,
                              font=("Courier New", 9))
        self._msg2.pack(pady=2)

        btn_f = tk.Frame(self, bg=WHITE)
        btn_f.pack(pady=10)
        tk.Button(btn_f, text="CHECK & CONFIRM",
                  bg=ACCENT, fg=WHITE,
                  font=("Courier New", 10, "bold"), relief='flat',
                  padx=14, pady=7, cursor='hand2',
                  command=self._check_and_preview).pack(side='left', padx=6)
        if idx > 0:
            tk.Button(btn_f, text="← BACK", bg=PANEL_BG, fg=TEXT,
                      font=("Courier New", 10), relief='flat',
                      padx=10, pady=7, cursor='hand2',
                      command=self._go_back).pack(side='left', padx=6)
        tk.Button(btn_f, text="CANCEL", bg=PANEL_BG, fg=TEXT,
                  font=("Courier New", 10), relief='flat',
                  padx=10, pady=7, cursor='hand2',
                  command=self.destroy).pack(side='left', padx=6)

    def _go_back(self):
        self._current_idx -= 1
        if self._added_units:
            self._added_units.pop()
        self._build_step2()

    def _check_and_preview(self):
        lab  = self._lab_var.get().strip()
        addr = self._addr_var.get().strip()

        if not lab:
            self._msg2.config(text="⚠  Lab / Log Number is required.")
            return
        if not addr:
            self._msg2.config(text="⚠  Address Code is required.")
            return

        key = f'CHP ({addr})'
        if key in self.app.area_result:
            ind = self.app.area_result.index(key)
            a1  = self.app.area_result[ind + 2]
            a2  = self.app.area_result[ind + 3]
            self._addr_display.config(text=f"  {a1}\n  {a2}", fg=GREEN)
            self._a1 = a1
            self._a2 = a2
        else:
            self._msg2.config(text=f"⚠  Address code '{addr}' not found.")
            self._addr_display.config(text="", )
            return

        self._msg2.config(text="")
        self._build_readback(lab, addr)

    def _build_readback(self, lab, addr):
        self._clear()
        utype  = self._unit_type
        idx    = self._current_idx
        total  = self._total

        tk.Label(self, text=f"CONFIRM ENTRY {idx + 1} OF {total}",
                 bg=WHITE, fg=ACCENT,
                 font=("Courier New", 12, "bold")).pack(pady=(16, 4))
        tk.Label(self, text="Read each digit carefully before confirming:",
                 bg=WHITE, fg=TEXT_DIM,
                 font=("Courier New", 9)).pack(pady=(0, 8))

        tk.Label(self, text="LAB / LOG NUMBER:", bg=WHITE, fg=TEXT_DIM,
                 font=("Courier New", 9, "bold")).pack(anchor='w', padx=20)
        self._render_digits(lab)

        tk.Label(self, text="ADDRESS CODE:", bg=WHITE, fg=TEXT_DIM,
                 font=("Courier New", 9, "bold")).pack(anchor='w', padx=20, pady=(8, 0))
        self._render_digits(addr)

        tk.Label(self, text=f"  {self._a1}\n  {self._a2}",
                 bg=WHITE, fg=GREEN,
                 font=("Courier New", 9)).pack(anchor='w', padx=20, pady=(2, 8))

        btn_f = tk.Frame(self, bg=WHITE)
        btn_f.pack(pady=8)
        tk.Button(btn_f, text="✓  CONFIRM",
                  bg=GREEN, fg=WHITE,
                  font=("Courier New", 10, "bold"), relief='flat',
                  padx=14, pady=7, cursor='hand2',
                  command=lambda: self._confirm_unit(lab, addr)).pack(side='left', padx=6)
        tk.Button(btn_f, text="✗  RE-ENTER",
                  bg=RED, fg=WHITE,
                  font=("Courier New", 10), relief='flat',
                  padx=10, pady=7, cursor='hand2',
                  command=self._build_step2).pack(side='left', padx=6)
        tk.Button(btn_f, text="CANCEL", bg=PANEL_BG, fg=TEXT,
                  font=("Courier New", 10), relief='flat',
                  padx=10, pady=7, cursor='hand2',
                  command=self.destroy).pack(side='left', padx=6)

    def _render_digits(self, value):
        positions = ['1ST','2ND','3RD','4TH','5TH','6TH',
                     '7TH','8TH','9TH','10TH','11TH','12TH']
        dframe = tk.Frame(self, bg=WHITE)
        dframe.pack(padx=20, pady=4)
        for i, ch in enumerate(value):
            box = tk.Frame(dframe, bg=PANEL_BG, bd=1, relief='solid',
                           width=40, height=52)
            box.pack_propagate(False)
            box.pack(side='left', padx=2)
            tk.Label(box, text=ch, bg=PANEL_BG, fg=ACCENT,
                     font=("Courier New", 18, "bold")).pack(expand=True)
            pos = positions[i] if i < len(positions) else f'{i+1}TH'
            tk.Label(box, text=pos, bg=PANEL_BG, fg=TEXT_DIM,
                     font=("Courier New", 6)).pack(pady=(0, 3))

    def _confirm_unit(self, lab, addr):
        unit = {
            'type':         self._unit_type,
            'lab_number':   lab,
            'address_code': addr,
            'address_1':    self._a1,
            'address_2':    self._a2,
            'unit_name':    self._unit_type,
            'serial':       '—',
            'chps':         '—',
            'manual':       True,
        }
        self._added_units.append(unit)
        self._current_idx += 1

        if self._current_idx < self._total:
            self._build_step2()
        else:
            self.destroy()
            self.on_complete(self._added_units)


# ─────────────────────────────────────────────
#  SHIPPING LABEL BUILDER — FIX 4: Single column layout
# ─────────────────────────────────────────────
def _build_shipping_doc(label_areas, units_by_area, date_str, total_units):
    """
    FIX 4: Single-column label layout — one label per row, full page width.
    Each label: area code, date, lab numbers, address. Bold 9pt Arial.
    Labels print on one side; shipping tracker goes on the reverse.
    """
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    BLACK = RGBColor(0x00, 0x00, 0x00)
    DARK  = RGBColor(0x1A, 0x1A, 0x1A)
    GREY  = RGBColor(0x55, 0x55, 0x55)
    SZ    = 10   # Slightly larger for single-column readability

    def _bg(cell, h):
        tc=cell._tc; tcPr=tc.get_or_add_tcPr()
        shd=OxmlElement('w:shd'); shd.set(qn('w:val'),'clear')
        shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),h)
        for o in tcPr.findall(qn('w:shd')): tcPr.remove(o)
        tcPr.append(shd)

    def _margins(cell, top=60, bottom=60, left=120, right=120):
        tc=cell._tc; tcPr=tc.get_or_add_tcPr(); mar=OxmlElement('w:tcMar')
        for s,v in [('top',top),('bottom',bottom),('left',left),('right',right)]:
            el=OxmlElement(f'w:{s}'); el.set(qn('w:w'),str(v)); el.set(qn('w:type'),'dxa')
            mar.append(el)
        for o in tcPr.findall(qn('w:tcMar')): tcPr.remove(o)
        tcPr.append(mar)

    def _tbl_borders(tbl, size=8, color='000000'):
        t=tbl._tbl; tblPr=t.find(qn('w:tblPr'))
        if tblPr is None: tblPr=OxmlElement('w:tblPr'); t.insert(0,tblPr)
        tblB=OxmlElement('w:tblBorders')
        for s in ('top','left','bottom','right','insideH','insideV'):
            el=OxmlElement(f'w:{s}')
            el.set(qn('w:val'),'single' if size>0 else 'none')
            el.set(qn('w:sz'),str(size)); el.set(qn('w:space'),'0')
            el.set(qn('w:color'),color); tblB.append(el)
        for o in tblPr.findall(qn('w:tblBorders')): tblPr.remove(o)
        tblPr.append(tblB)

    def _tbl_width(tbl, w):
        t=tbl._tbl; tblPr=t.find(qn('w:tblPr'))
        if tblPr is None: tblPr=OxmlElement('w:tblPr'); t.insert(0,tblPr)
        tblW=OxmlElement('w:tblW'); tblW.set(qn('w:w'),str(w))
        tblW.set(qn('w:type'),'dxa')
        for o in tblPr.findall(qn('w:tblW')): tblPr.remove(o)
        tblPr.append(tblW)

    def _zero_spacing(para):
        pPr=para._p.get_or_add_pPr(); sp=OxmlElement('w:spacing')
        sp.set(qn('w:before'),'0'); sp.set(qn('w:after'),'0')
        for o in pPr.findall(qn('w:spacing')): pPr.remove(o)
        pPr.append(sp)

    def _run(para, text, bold=False, color=DARK, size=SZ):
        para.clear(); _zero_spacing(para)
        r=para.add_run(text); r.bold=bold; r.font.size=Pt(size)
        r.font.color.rgb=color; r.font.name='Arial'

    def _add_line(cell, text, bold=False, color=DARK, size=SZ):
        p=cell.add_paragraph(); _zero_spacing(p)
        r=p.add_run(text); r.bold=bold; r.font.size=Pt(size)
        r.font.color.rgb=color; r.font.name='Arial'

    # FIX 4: Full page width for single column
    # Letter: 8.5" = 12240 DXA, margins 0.5" each side → usable = 7.5" = 10800 DXA
    PAGE_W = 10800

    def build_label_content(cell, area_code, lab_numbers, unit_names,
                             address_1, address_2, unit_count, date_str):
        """
        Build label content directly into the given table cell.
        FIX 4: Uses address_code correctly as 'CHP (area_code)'.
        """
        _margins(cell, top=80, bottom=80, left=160, right=160)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        # Clear default paragraph
        for p in cell.paragraphs:
            for r in p.runs:
                r.text = ''

        # Line 1: Date + unit count (header line)
        _run(cell.paragraphs[0],
             f"# {unit_count} UNIT{'S' if unit_count != 1 else ''}    {date_str}",
             bold=True, color=BLACK, size=SZ)

        # Lines: each lab number with unit model
        for lab, uname in zip(lab_numbers, unit_names):
            display = f"{lab}    {uname}" if uname and uname != lab else lab
            _add_line(cell, display, bold=True, color=BLACK, size=SZ)

        # FIX 4: Print the FULL CHP area code string properly
        _add_line(cell, f"CHP ({area_code})", bold=True, color=BLACK, size=SZ)

        # Standard recipient title
        _add_line(cell, "Traffic Radar Coordinator", bold=False, color=DARK, size=SZ)

        # Address lines — print each part on its own line for clarity
        if address_1:
            _add_line(cell, address_1, bold=False, color=DARK, size=SZ)
        if address_2:
            _add_line(cell, address_2, bold=False, color=DARK, size=SZ)

    # ── Build document ────────────────────────────────────────────────────────
    doc = docx.Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(0.5)
        sec.bottom_margin = Inches(0.5)
        sec.left_margin   = Inches(0.5)
        sec.right_margin  = Inches(0.5)
    # Remove default empty paragraph
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    def _zero_spacing_para(para):
        pPr=para._p.get_or_add_pPr(); sp=OxmlElement('w:spacing')
        sp.set(qn('w:before'),'0'); sp.set(qn('w:after'),'60')
        for o in pPr.findall(qn('w:spacing')): pPr.remove(o)
        pPr.append(sp)

    # FIX 4: Single column — one label per row in a 1-column table
    tbl = doc.add_table(rows=0, cols=1)
    tbl.style = 'Table Grid'
    _tbl_borders(tbl, size=8)
    _tbl_width(tbl, PAGE_W)

    for area in label_areas:
        units  = units_by_area[area]
        row    = tbl.add_row()
        cell   = row.cells[0]

        # Set row height (auto — content determines height)
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _OE
        trPr = row._tr.get_or_add_trPr()
        trH  = _OE('w:trHeight')
        trH.set(_qn('w:val'), '1200')   # minimum height = ~1 cm
        trH.set(_qn('w:hRule'), 'atLeast')
        trPr.append(trH)

        build_label_content(
            cell,
            area_code  = units[0].get('address_code', area),
            lab_numbers= [u['lab_number'] for u in units],
            unit_names = [u.get('unit_name', '') for u in units],
            address_1  = units[0].get('address_1', ''),
            address_2  = units[0].get('address_2', ''),
            unit_count = len(units),
            date_str   = date_str,
        )

    return doc


# ─────────────────────────────────────────────
#  GROUPED LABEL BUILDER — FIX 5: NO CASE = separate label; HAS CASE = grouped
# ─────────────────────────────────────────────
def _build_shipping_doc_grouped(label_keys, units_by_label, date_str, total_units):
    """
    FIX 5: Renamed 'SEPARATE CASE' → 'NO CASE'.
    - NO CASE units: each gets its OWN separate label regardless of area code.
    - HAS CASE units: grouped together per area code onto one shared label.
    Delegates final rendering to _build_shipping_doc (single column).
    """
    compatible_units_by_area = {}
    for key in label_keys:
        group = units_by_label[key]
        all_cased = all(u.get('_has_case', False) for u in group)
        tagged = []
        for u in group:
            uc = dict(u)
            if not all_cased:
                # FIX 5: Mark as [NO CASE] instead of [WITH CASE]
                uc['unit_name'] = uc.get('unit_name', '') + '  [NO CASE]'
            # cased units: no tag, they group naturally
            tagged.append(uc)
        compatible_units_by_area[key] = tagged

    return _build_shipping_doc(label_keys, compatible_units_by_area, date_str, total_units)


# ─────────────────────────────────────────────
#  OLD CERT RESOLVER
# ─────────────────────────────────────────────
def _resolve_old_cert(log_num, app):
    log_num = log_num.strip()
    if not log_num:
        return {'ok': False, 'error': 'Empty log number'}

    if log_num.upper().startswith('ASL'):
        unit_type = 'LIDAR'
        cert_path = os.path.join(LIDAR_OUT_DIR, f"{log_num}.docx")
        sheet_name = 'LIDAR'
        prefix = LIDAR_PREFIX
    elif log_num.upper().startswith('AS'):
        unit_type = 'RADAR'
        cert_path = os.path.join(RADAR_OUT_DIR, f"{log_num}.docx")
        sheet_name = 'RADAR'
        prefix = RADAR_PREFIX
    else:
        return {'ok': False, 'error': f"Unknown prefix — expected AS26- or ASL26-"}

    if not os.path.exists(cert_path):
        return {'ok': False, 'error': f"Cert file not found: {os.path.basename(cert_path)}"}

    address_code = ''
    address_1    = ''
    address_2    = ''
    serial       = '—'
    chps         = '—'
    unit_name    = unit_type

    try:
        wb = load_workbook(EXCEL_2026, read_only=True, data_only=True)
        if sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            for row_tuple in ws.iter_rows(min_row=2, values_only=True):
                if not row_tuple or not row_tuple[0]:
                    continue
                lab_cell = str(row_tuple[0]).replace(' (RETEST)', '').strip()
                if lab_cell.upper() == log_num.upper():
                    address_code = str(row_tuple[3]) if row_tuple[3] else ''
                    chps_raw     = str(row_tuple[2]) if row_tuple[2] else ''
                    chps         = chps_raw.replace('CHPS', '')
                    serial_raw   = str(row_tuple[1]) if row_tuple[1] else ''
                    for pfx in ['DS','DH','DE','AS','ZC','ZM','DC',
                                  'TS','TJ_SXB','TJ_S','UX','LP','UL']:
                        if serial_raw.startswith(pfx):
                            serial = serial_raw[len(pfx):]
                            break
                    else:
                        serial = serial_raw
                    unit_name = str(row_tuple[8]) if len(row_tuple) > 8 and row_tuple[8] else unit_type
                    break
        wb.close()
    except Exception:
        pass

    if not address_code:
        pattern = os.path.join(HISTORY_DIR, 'week_report_*.xlsx')
        history_files = sorted(glob.glob(pattern), reverse=True)
        for filepath in history_files:
            try:
                wb = load_workbook(filepath, read_only=True, data_only=True)
                if sheet_name not in wb.sheetnames:
                    wb.close(); continue
                ws = wb[sheet_name]
                for row_tuple in ws.iter_rows(min_row=2, values_only=True):
                    if not row_tuple or not row_tuple[0]:
                        continue
                    lab_cell = str(row_tuple[0]).replace(' (RETEST)', '').strip()
                    if lab_cell.upper() == log_num.upper():
                        address_code = str(row_tuple[3]) if row_tuple[3] else ''
                        chps_raw     = str(row_tuple[2]) if row_tuple[2] else ''
                        chps         = chps_raw.replace('CHPS', '')
                        serial_raw   = str(row_tuple[1]) if row_tuple[1] else ''
                        for pfx in ['DS','DH','DE','AS','ZC','ZM','DC',
                                      'TS','TJ_SXB','TJ_S','UX','LP','UL']:
                            if serial_raw.startswith(pfx):
                                serial = serial_raw[len(pfx):]
                                break
                        else:
                            serial = serial_raw
                        unit_name = str(row_tuple[8]) if len(row_tuple) > 8 and row_tuple[8] else unit_type
                        wb.close()
                        break
                else:
                    wb.close()
                    continue
                break
            except Exception:
                continue

    if address_code and hasattr(app, 'area_result'):
        key = f'CHP ({address_code})'
        if key in app.area_result:
            ind = app.area_result.index(key)
            address_1 = app.area_result[ind + 2]
            address_2 = app.area_result[ind + 3]

    unit = {
        'type':         unit_type,
        'lab_number':   log_num,
        'address_code': address_code or '—',
        'address_1':    address_1,
        'address_2':    address_2,
        'unit_name':    unit_name,
        'serial':       serial,
        'chps':         chps,
        'loaded':       True,
    }
    return {'ok': True, 'unit': unit}


# ─────────────────────────────────────────────
#  SHIPPING TAB — FIX 5: "NO CASE" replaces "SEPARATE CASE"
# ─────────────────────────────────────────────
class ShippingTab(tk.Frame):
    def __init__(self, parent, app):
        super().__init__(parent, bg=BG)
        self.app = app
        self._fail_vars    = {}
        self._has_case_vars = {}
        self._build()

    def _build(self):
        hdr = tk.Frame(self, bg=BG)
        hdr.pack(fill='x', pady=(4, 2))
        tk.Label(hdr, text="\u25b8 SESSION UNITS — SHIPPING",
                 bg=BG, fg=ACCENT,
                 font=("Courier New", 10, "bold")).pack(side='left')
        tk.Frame(hdr, bg=ACCENT, height=1).pack(
            side='left', fill='x', expand=True, padx=8)

        legend = tk.Frame(self, bg=BG)
        legend.pack(fill='x', pady=(0, 6))
        tk.Label(legend,
                 text="  \u2610 Unchecked = WILL SHIP    \u2611 Checked = FAILED (excluded)"
                      "      NO CASE checkbox = separate label per unit",
                 bg=BG, fg=TEXT_DIM,
                 font=("Courier New", 9)).pack(side='left')

        # FIX 5: Column header renamed from "SEPERATE LABEL" to "NO CASE"
        col_hdr = tk.Frame(self, bg=CARD_BG)
        col_hdr.pack(fill='x', pady=(0, 2))
        for col, w in [("FAIL?", 6), ("LAB NUMBER", 16), ("TYPE", 6),
                        ("SERIAL", 10), ("CHPS", 8), ("ADDR", 6), ("UNIT", 18), ("NO CASE", 8)]:
            tk.Label(col_hdr, text=col, bg=CARD_BG, fg=ACCENT,
                     font=("Courier New", 8, "bold"),
                     width=w, anchor='w').pack(side='left', padx=4, pady=4)

        list_frame = tk.Frame(self, bg=BG)
        list_frame.pack(fill='both', expand=True)

        self.list_canvas = tk.Canvas(list_frame, bg=BG, highlightthickness=0)
        sb = ttk.Scrollbar(list_frame, orient='vertical',
                           command=self.list_canvas.yview)
        self.list_canvas.configure(yscrollcommand=sb.set)
        sb.pack(side='right', fill='y')
        self.list_canvas.pack(side='left', fill='both', expand=True)

        self.list_inner = tk.Frame(self.list_canvas, bg=BG)
        self._list_win = self.list_canvas.create_window(
            (0, 0), window=self.list_inner, anchor='nw')

        self.list_inner.bind('<Configure>',
            lambda e: self.list_canvas.configure(
                scrollregion=self.list_canvas.bbox('all')))
        self.list_canvas.bind('<Configure>',
            lambda e: self.list_canvas.itemconfig(self._list_win, width=e.width))

        bottom = tk.Frame(self, bg=PANEL_BG)
        bottom.pack(fill='x', side='bottom')

        row1 = tk.Frame(bottom, bg=PANEL_BG)
        row1.pack(fill='x', padx=12, pady=(6, 0))
        self.summary_lbl = tk.Label(
            row1, text="No units in this session yet.",
            bg=PANEL_BG, fg=TEXT_DIM,
            font=("Courier New", 10))
        self.summary_lbl.pack(side='left')

        row2 = tk.Frame(bottom, bg=PANEL_BG)
        row2.pack(fill='x', padx=12, pady=(2, 8))

        tk.Button(row2, text="+ ADD",
            bg=GREEN, fg=WHITE,
            font=("Courier New", 9, "bold"), relief='flat',
            padx=10, pady=6, cursor='hand2',
            command=self._open_add_dialog).pack(side='left', padx=(0, 6))

        tk.Button(row2, text="REFRESH",
            bg=PANEL_BG, fg=TEXT_DIM,
            font=("Courier New", 9), relief='flat',
            padx=10, pady=6, cursor='hand2',
            command=self.refresh).pack(side='left', padx=(0, 6))

        tk.Button(row2, text="CLEAR SESSION",
            bg=PANEL_BG, fg=RED,
            font=("Courier New", 9, "bold"), relief='flat',
            padx=10, pady=6, cursor='hand2',
            command=self._clear_session).pack(side='left')

        tk.Button(row2, text="LOAD OLD CERT",
            bg=PANEL_BG, fg=ACCENT,
            font=("Courier New", 9, "bold"), relief='flat',
            padx=10, pady=6, cursor='hand2',
            command=self._load_old_cert).pack(side='left', padx=(6, 0))

        self.gen_btn = tk.Button(row2,
            text="  GENERATE SHIPPING DOC  ",
            bg=ACCENT, fg=WHITE,
            font=("Courier New", 10, "bold"), relief='flat',
            padx=14, pady=6, cursor='hand2',
            command=self._generate)
        self.gen_btn.pack(side='right', padx=(6, 0))

        # FIX 5: Renamed button from "SEPARATE LABELS" to "ALL SEPARATE"
        self.sep_lbl_btn = tk.Button(row2,
            text="ALL SEPARATE",
            bg=PANEL_BG, fg=ACCENT,
            font=("Courier New", 9, "bold"), relief='flat',
            padx=10, pady=6, cursor='hand2',
            command=lambda: self._generate(force_separate=True))
        self.sep_lbl_btn.pack(side='right', padx=(0, 4))

        self.seal_btn = tk.Button(row2,
            text="  SEAL CERTS  ",
            bg=AMBER, fg=WHITE,
            font=("Courier New", 10, "bold"), relief='flat',
            padx=14, pady=6, cursor='hand2',
            command=self._seal_certs)
        self.seal_btn.pack(side='right', padx=(0, 6))

    def refresh(self):
        for w in self.list_inner.winfo_children():
            w.destroy()

        units = self.app.session_units
        if not units:
            tk.Label(self.list_inner,
                     text="\n  No units entered yet this session.\n  Enter units in the RADAR or LIDAR tab.",
                     bg=BG, fg=TEXT_DIM,
                     font=("Courier New", 10), justify='left').pack(pady=20)
            self.summary_lbl.config(text="No units in this session yet.")
            return

        existing_keys = set(self._fail_vars.keys())
        for u in units:
            key = u.get('lab_number', '')
            if key not in self._fail_vars:
                self._fail_vars[key] = tk.BooleanVar(value=False)
            if key not in self._has_case_vars:
                utype = u.get('type', '')
                unit_name = u.get('unit_name', '')
                default_case = utype == 'RADAR' and unit_name in ('Stalker II SDR',) or \
                               'Ultralyte' in unit_name
                self._has_case_vars[key] = tk.BooleanVar(value=default_case)

        current_keys = {u.get('lab_number', '') for u in units}
        for k in list(self._fail_vars.keys()):
            if k not in current_keys:
                del self._fail_vars[k]
        for k in list(self._has_case_vars.keys()):
            if k not in current_keys:
                del self._has_case_vars[k]

        for i, u in enumerate(units):
            key      = u['lab_number']
            fail_var = self._fail_vars[key]
            row_bg   = BG if i % 2 == 0 else PANEL_BG

            row = tk.Frame(self.list_inner, bg=row_bg)
            row.pack(fill='x', pady=1)

            cb = tk.Checkbutton(
                row, variable=fail_var,
                bg=row_bg, activebackground=row_bg,
                selectcolor=PANEL_BG,
                command=lambda r=row, v=fail_var: self._on_toggle(r, v))
            cb.pack(side='left', padx=6)

            for val, w in [
                (u.get('lab_number','—'),    16),
                (u.get('type','—'),            6),
                (u.get('serial',''),          10),
                (u.get('chps',''),             8),
                (u.get('address_code','—'),    6),
                (u.get('unit_name','—'),      18),
            ]:
                lbl = tk.Label(row, text=str(val), bg=row_bg,
                               fg=TEXT_DIM if fail_var.get() else TEXT,
                               font=("Courier New", 9),
                               width=w, anchor='w')
                lbl.pack(side='left', padx=4, pady=4)

            # FIX 5: "NO CASE" checkbox — checked = no case = separate label
            no_case_var = self._has_case_vars[key]
            # Invert: the column is "NO CASE", so we store has_case=True meaning grouped
            # Display: checked = NO CASE (separate), unchecked = HAS CASE (group)
            # We keep the storage as _has_case_vars (True=has case=group),
            # but visually the checkbox is "NO CASE" — checked means no case
            no_case_display_var = tk.BooleanVar(value=not no_case_var.get())

            def _make_nocase_toggle(ncdv, hcv):
                def _toggle():
                    new_nocase = ncdv.get()
                    hcv.set(not new_nocase)
                return _toggle

            no_case_cb = tk.Checkbutton(
                row, variable=no_case_display_var,
                bg=row_bg, activebackground=row_bg,
                selectcolor=PANEL_BG,
                command=_make_nocase_toggle(no_case_display_var, no_case_var))
            no_case_cb.pack(side='left', padx=2)
            # Store the display var so refresh doesn't lose state
            self._has_case_vars[key]._nocase_display = no_case_display_var

            status = tk.Label(row, text="FAILED" if fail_var.get() else "",
                              bg=row_bg, fg=RED,
                              font=("Courier New", 8, "bold"), width=8)
            status.pack(side='left', padx=4)

        total  = len(units)
        failed = sum(1 for u in units if self._fail_vars[u['lab_number']].get())
        ship   = total - failed
        self.summary_lbl.config(
            text=f"Total: {total}   \u2713 To ship: {ship}   \u2717 Failed: {failed}",
            fg=TEXT)

    def _on_toggle(self, row, fail_var):
        self.refresh()

    def _seal_certs(self):
        units = self.app.session_units
        if not units:
            messagebox.showwarning("No Units", "No units in this session to seal.")
            return

        if not os.path.exists(SEAL_IMAGE):
            messagebox.showerror(
                "Seal Image Missing",
                f"Seal image not found at:\n{SEAL_IMAGE}\n\n"
                f"Please place your seal PNG file at that path and try again.")
            return

        to_seal = [u for u in units
                   if not self._fail_vars.get(u['lab_number'],
                                               tk.BooleanVar()).get()]
        if not to_seal:
            messagebox.showwarning("Nothing to Seal",
                                   "All units are marked as failed.")
            return

        date_dlg = tk.Toplevel(self.winfo_toplevel())
        date_dlg.title("Date Certified")
        date_dlg.configure(bg=WHITE)
        date_dlg.resizable(False, False)
        date_dlg.grab_set()
        date_dlg.update_idletasks()
        sw = date_dlg.winfo_screenwidth()
        sh = date_dlg.winfo_screenheight()
        date_dlg.geometry(f"380x230+{(sw-380)//2}+{(sh-230)//2}")

        tk.Label(date_dlg, text="DATE CERTIFIED",
                 bg=WHITE, fg=ACCENT,
                 font=("Courier New", 12, "bold")).pack(pady=(20, 4))
        tk.Label(date_dlg,
                 text=f"Stamped on all {len(to_seal)} cert(s). Edit if needed.",
                 bg=WHITE, fg=TEXT_DIM,
                 font=("Courier New", 9)).pack(pady=(0, 10))

        date_var = tk.StringVar(value=datetime.now().strftime("%m/%d/%Y"))
        date_entry = tk.Entry(date_dlg, textvariable=date_var,
                              bg=PANEL_BG, fg=TEXT, insertbackground=TEXT,
                              font=("Courier New", 14, "bold"),
                              width=14, relief='solid', bd=1, justify='center')
        date_entry.pack(ipady=6)
        date_entry.select_range(0, 'end')
        date_entry.focus_set()

        date_result = [None]

        def on_date_confirm():
            date_result[0] = date_var.get().strip()
            date_dlg.destroy()
        def on_date_cancel():
            date_dlg.destroy()

        btn_f = tk.Frame(date_dlg, bg=WHITE)
        btn_f.pack(pady=14)
        cb = tk.Button(btn_f, text="PROCEED TO SEAL",
                  bg=ACCENT, fg=WHITE,
                  font=("Courier New", 10, "bold"), relief='flat',
                  padx=14, pady=7, cursor='hand2', command=on_date_confirm)
        cb.pack(side='left', padx=6)
        cb.bind("<Return>", lambda e: on_date_confirm())
        date_entry.bind("<Return>", lambda e: on_date_confirm())
        tk.Button(btn_f, text="CANCEL", bg=PANEL_BG, fg=TEXT,
                  font=("Courier New", 10), relief='flat',
                  padx=14, pady=7, cursor='hand2',
                  command=on_date_cancel).pack(side='left', padx=6)
        date_dlg.protocol("WM_DELETE_WINDOW", on_date_cancel)
        self.winfo_toplevel().wait_window(date_dlg)

        if date_result[0] is None:
            return

        cert_date = date_result[0]
        os.makedirs(SEALED_DIR, exist_ok=True)

        success_list = []
        fail_list    = []
        warn_list    = []

        for u in to_seal:
            lab  = u['lab_number']
            utype = u['type']

            if utype == 'RADAR':
                src_path = os.path.join(RADAR_OUT_DIR, f"{lab}.docx")
            else:
                src_path = os.path.join(LIDAR_OUT_DIR, f"{lab}.docx")

            dst_path = os.path.join(SEALED_DIR, f"{lab}_sealed.docx")

            ok, msg = seal_document(src_path, dst_path, utype, cert_date)
            if ok:
                if msg:
                    warn_list.append(f"{lab}: {msg}")
                else:
                    success_list.append(lab)
            else:
                fail_list.append(f"{lab}: {msg}")

        lines = []
        if success_list:
            lines.append(f"✓ Sealed successfully ({len(success_list)}):")
            for lab in success_list:
                lines.append(f"   {lab}_sealed.docx")
        if warn_list:
            lines.append(f"\n⚠ Sealed with warnings ({len(warn_list)}):")
            for w in warn_list:
                lines.append(f"   {w}")
        if fail_list:
            lines.append(f"\n✗ Failed ({len(fail_list)}):")
            for f_ in fail_list:
                lines.append(f"   {f_}")

        lines.append(f"\nSaved to:\n  {SEALED_DIR}")

        if fail_list:
            messagebox.showwarning("Sealing Complete — Some Errors",
                                   "\n".join(lines))
        else:
            messagebox.showinfo("Sealing Complete", "\n".join(lines))

    def _load_old_cert(self):
        dlg = tk.Toplevel(self.winfo_toplevel())
        dlg.title("Load Old Certificate(s)")
        dlg.configure(bg=WHITE)
        dlg.resizable(False, False)
        dlg.grab_set()
        dlg.update_idletasks()
        sw = dlg.winfo_screenwidth(); sh = dlg.winfo_screenheight()
        dlg.geometry(f"480x460+{(sw-480)//2}+{(sh-460)//2}")

        tk.Label(dlg, text="LOAD OLD CERTIFICATE(S)",
                 bg=WHITE, fg=ACCENT,
                 font=("Courier New", 12, "bold")).pack(pady=(18, 2))
        tk.Label(dlg,
                 text="Enter log numbers or ranges to add to the shipping list for sealing.\n"
                      "Separate multiple entries with commas or new lines.",
                 bg=WHITE, fg=TEXT_DIM,
                 font=("Courier New", 9), justify='center').pack(pady=(0, 10))

        txt_frame = tk.Frame(dlg, bg=WHITE)
        txt_frame.pack(fill='x', padx=20)
        tk.Label(txt_frame, text="Log Number(s):",
                 bg=WHITE, fg=TEXT,
                 font=("Courier New", 10), anchor='w').pack(anchor='w')
        txt = tk.Text(txt_frame, bg=PANEL_BG, fg=TEXT, insertbackground=TEXT,
                      font=("Courier New", 12, "bold"),
                      width=36, height=6, relief='solid', bd=1)
        txt.pack(fill='x', ipady=4)
        txt.focus_set()
        tk.Label(txt_frame,
                 text="e.g.  AS26-72\nASL26-45, ASL26-46\nAS26-100 to AS26-108",
                 bg=WHITE, fg=TEXT_DIM,
                 font=("Courier New", 8), justify='left').pack(anchor='w', pady=(2,0))

        preview_frame = tk.Frame(dlg, bg=PANEL_BG, relief='solid', bd=1)
        preview_frame.pack(fill='x', padx=20, pady=(10, 0))
        self._load_preview = tk.Label(preview_frame, text="",
                                       bg=PANEL_BG, fg=TEXT,
                                       font=("Courier New", 8),
                                       justify='left', anchor='w',
                                       wraplength=420)
        self._load_preview.pack(fill='x', padx=8, pady=6)

        msg_lbl = tk.Label(dlg, text="", bg=WHITE, fg=RED,
                           font=("Courier New", 9))
        msg_lbl.pack(pady=(4,0))

        btn_f = tk.Frame(dlg, bg=WHITE)
        btn_f.pack(pady=12)

        def on_load():
            raw = txt.get("1.0", "end").strip()
            if not raw:
                msg_lbl.config(text="⚠  Enter at least one log number.")
                return
            entries, parse_errors = expand_log_number_entries(raw)
            if not entries:
                msg_lbl.config(text="⚠  No valid log numbers found.")
                if parse_errors:
                    self._load_preview.config(text="\n".join(parse_errors), fg=RED)
                return

            added = []
            errors = list(parse_errors)
            for log_num in entries:
                result = _resolve_old_cert(log_num, self.app)
                if result['ok']:
                    added.append(result['unit'])
                else:
                    errors.append(f"{log_num}: {result['error']}")

            lines = []
            if added:
                lines.append(f"✓ Ready to add ({len(added)}):")
                for u in added:
                    lines.append(f"  {u['lab_number']}  {u['type']}  addr:{u['address_code']}")
            if errors:
                lines.append(f"\n✗ Errors ({len(errors)}):")
                for e in errors:
                    lines.append(f"  {e}")
            self._load_preview.config(
                text="\n".join(lines),
                fg=GREEN if not errors else AMBER)

            if not added:
                msg_lbl.config(text="⚠  No certs could be resolved.")
                return

            for u in added:
                key = u['lab_number']
                if key not in [x['lab_number'] for x in self.app.session_units]:
                    self.app.session_units.append(u)
            save_session(self.app.session_units)
            count = len(self.app.session_units)
            self.app.shipping_tab_btn.config(text=f"  SHIPPING ({count})  ")
            self.refresh()
            dlg.destroy()
            messagebox.showinfo("Loaded",
                f"{len(added)} cert(s) added to shipping list.\n" +
                (f"{len(errors)} could not be resolved." if errors else ""))

        ok_btn = tk.Button(btn_f, text="LOAD INTO SESSION",
                  bg=ACCENT, fg=WHITE,
                  font=("Courier New", 10, "bold"), relief='flat',
                  padx=14, pady=7, cursor='hand2',
                  command=on_load)
        ok_btn.pack(side='left', padx=6)
        ok_btn.bind("<Return>", lambda e: on_load())
        txt.bind("<Control-Return>", lambda e: on_load())

        tk.Button(btn_f, text="CANCEL", bg=PANEL_BG, fg=TEXT,
                  font=("Courier New", 10), relief='flat',
                  padx=14, pady=7, cursor='hand2',
                  command=dlg.destroy).pack(side='left', padx=6)

    def _open_add_dialog(self):
        def on_complete(units):
            for u in units:
                self.app.session_units.append(u)
            save_session(self.app.session_units)
            count = len(self.app.session_units)
            self.app.shipping_tab_btn.config(
                text=f"  SHIPPING ({count})  ")
            self.refresh()
            messagebox.showinfo(
                "Units Added",
                f"{len(units)} unit(s) added to shipping list.")
        ManualAddDialog(self.winfo_toplevel(), self.app, on_complete)

    def _clear_session(self):
        if not self.app.session_units:
            messagebox.showinfo("Clear Session", "Session is already empty.")
            return
        ans = messagebox.askyesno(
            "Clear Session",
            f"This will remove all {len(self.app.session_units)} unit(s) "
            f"from the shipping list.\n\nThis does NOT delete any certificates "
            f"or Excel entries.\n\nAre you sure?")
        if not ans:
            return
        self.app.session_units.clear()
        self._fail_vars.clear()
        clear_session_file()
        self.app.shipping_tab_btn.config(text="  SHIPPING  ")
        self.refresh()
        messagebox.showinfo("Cleared", "Session cleared.")

    def _generate(self, force_separate=False):
        units = self.app.session_units
        if not units:
            messagebox.showwarning("No Units", "No units in this session to ship.")
            return

        to_ship = [u for u in units
                   if not self._fail_vars.get(
                       u.get('lab_number', ''), tk.BooleanVar()).get()]

        if not to_ship:
            messagebox.showwarning("No Units to Ship",
                                   "All units are marked as failed.\nNothing to generate.")
            return

        # Snapshot the case flag
        for u in to_ship:
            key = u.get('lab_number', '')
            u['_has_case'] = self._has_case_vars.get(key, tk.BooleanVar()).get()

        failed_count = len(units) - len(to_ship)
        mode_str = "ALL SEPARATE labels" if force_separate else "NO CASE=separate, HAS CASE=grouped by area"
        ans = messagebox.askyesno(
            "Generate Shipping Document",
            f"Ready to generate shipping document:\n\n"
            f"  Units to ship:   {len(to_ship)}\n"
            f"  Excluded (fail): {failed_count}\n"
            f"  Label mode:      {mode_str}\n\n"
            f"Proceed?")
        if not ans:
            return

        EXCLUDED_AREA_CODES = {'645', '', None}

        # ── FIX 5: Grouping logic ────────────────────────────────────────────
        # NO CASE (has_case=False): separate label per unit regardless of area
        # HAS CASE (has_case=True): group by area code — one label per area
        label_order    = []
        units_by_label = {}

        def add_to_group(key, u):
            if key not in units_by_label:
                label_order.append(key)
                units_by_label[key] = []
            units_by_label[key].append(u)

        skipped_areas = set()
        for u in to_ship:
            area     = str(u.get('address_code', '') or '').strip()
            lab      = u.get('lab_number', '')
            has_case = u['_has_case']

            if area in EXCLUDED_AREA_CODES or area == '645':
                skipped_areas.add(area)
                add_to_group(f'_excl_{area}_{lab}', u)
                continue

            if force_separate:
                # All separate
                add_to_group(f'sep_{lab}', u)
            elif not has_case:
                # FIX 5: NO CASE → SEPARATE label per unit
                add_to_group(f'nocase_{lab}', u)
            else:
                # FIX 5: HAS CASE → GROUPED by area code
                add_to_group(f'area_{area}', u)

        label_keys_for_print = [k for k in label_order if not k.startswith('_excl_')]
        units_by_area_for_print = {k: units_by_label[k] for k in label_keys_for_print}

        today         = datetime.now()
        date_str      = today.strftime("%m/%d/%Y")
        file_date_str = today.strftime("%m%d%Y")
        total_units   = len(to_ship)
        suffix        = '_separate' if force_separate else ''
        out_filename  = f"shipping_{file_date_str}{suffix}.docx"
        out_path      = os.path.join(BASE_DIR, out_filename)

        try:
            doc = _build_shipping_doc_grouped(
                label_keys_for_print, units_by_area_for_print,
                date_str, total_units)
            doc.save(out_path)

            # Update column F (shipped date) in Excel
            excel_errors = []
            try:
                wb_xl = load_workbook(EXCEL_2026)
                for sheet_name_xl in ('RADAR', 'LIDAR'):
                    if sheet_name_xl not in wb_xl.sheetnames:
                        continue
                    ws_xl = wb_xl[sheet_name_xl]
                    for u in to_ship:
                        lab = str(u['lab_number']).replace(' (RETEST)', '').strip()
                        for row_idx in range(2, ws_xl.max_row + 1):
                            cell_a = ws_xl.cell(row=row_idx, column=1).value
                            if cell_a and str(cell_a).replace(' (RETEST)', '').strip().upper() == lab.upper():
                                ws_xl.cell(row=row_idx, column=6).value = date_str
                                break
                wb_xl.save(EXCEL_2026)
                wb_xl.close()
                try:
                    self.app.wb = load_workbook(EXCEL_2026)
                except Exception:
                    pass
            except Exception as e:
                excel_errors.append(str(e))

            msg = (f"Saved as:\n  {out_filename}\n\n"
                   f"Units shipped:    {total_units}\n"
                   f"Labels generated: {len(label_keys_for_print)}")
            if skipped_areas:
                msg += (f"\nNo label (excluded area codes): "
                        f"{', '.join(str(s) for s in skipped_areas)}")
            if excel_errors:
                msg += f"\n\n⚠ Excel update error:\n{excel_errors[0]}"
            messagebox.showinfo("Shipping Document Generated", msg)
            self.refresh()

        except Exception as e:
            messagebox.showerror("Error",
                                 f"Failed to generate shipping document:\n{e}")


# ─────────────────────────────────────────────
#  SEARCH TAB — FIX 6: Fetch ALL entries from ALL years
# ─────────────────────────────────────────────
class SearchTab(tk.Frame):
    def __init__(self, parent, app):
        super().__init__(parent, bg=BG)
        self.app = app
        self._build()

    def _build(self):
        hdr = tk.Frame(self, bg=BG)
        hdr.pack(fill='x', pady=(4, 2))
        tk.Label(hdr, text="\u25b8 UNIT HISTORY SEARCH",
                 bg=BG, fg=ACCENT,
                 font=("Courier New", 10, "bold")).pack(side='left')
        tk.Frame(hdr, bg=ACCENT, height=1).pack(
            side='left', fill='x', expand=True, padx=8)

        form = tk.Frame(self, bg=PANEL_BG, relief='flat')
        form.pack(fill='x', pady=(0, 8), ipady=10)

        row0 = tk.Frame(form, bg=PANEL_BG)
        row0.pack(fill='x', padx=16, pady=(10, 4))
        tk.Label(row0, text="Number of units to search:",
                 bg=PANEL_BG, fg=TEXT,
                 font=("Courier New", 10), width=26, anchor='w').pack(side='left')
        self._num_var = tk.StringVar(value='1')
        num_entry = tk.Entry(row0, textvariable=self._num_var,
                             bg=ENTRY_BG, fg=TEXT, insertbackground=TEXT,
                             font=("Courier New", 12, "bold"),
                             width=5, relief='solid', bd=1)
        num_entry.pack(side='left', ipady=4)
        num_entry.bind("<Return>", lambda e: self._type_combo.focus_set())
        self._num_entry = num_entry

        row1 = tk.Frame(form, bg=PANEL_BG)
        row1.pack(fill='x', padx=16, pady=4)

        tk.Label(row1, text="Unit Type:",
                 bg=PANEL_BG, fg=TEXT,
                 font=("Courier New", 10), width=12, anchor='w').pack(side='left')
        all_types = [("RADAR", t) for t in RADAR_TYPES] + [("LIDAR", t) for t in LIDAR_TYPES]
        type_labels = [f"{cat}: {t}" for cat, t in all_types]
        self._type_map  = {lbl: (cat, t) for lbl, (cat, t) in zip(type_labels, all_types)}
        self._type_var  = tk.StringVar(value=type_labels[0])
        self._type_combo = ttk.Combobox(row1, textvariable=self._type_var,
                                         values=type_labels, state='readonly',
                                         width=18, font=("Courier New", 10))
        self._type_combo.pack(side='left', padx=(0, 20))
        self._type_combo.bind("<Return>", lambda e: self._serial_entry.focus_set())

        tk.Label(row1, text="Serial Number:",
                 bg=PANEL_BG, fg=TEXT,
                 font=("Courier New", 10), width=14, anchor='w').pack(side='left')
        self._serial_var = tk.StringVar()
        self._serial_entry = tk.Entry(row1, textvariable=self._serial_var,
                                       bg=ENTRY_BG, fg=TEXT, insertbackground=TEXT,
                                       font=("Courier New", 12, "bold"),
                                       width=14, relief='solid', bd=1)
        self._serial_entry.pack(side='left', ipady=4)
        self._serial_entry.bind("<Return>", lambda e: self._do_search())

        row2 = tk.Frame(form, bg=PANEL_BG)
        row2.pack(fill='x', padx=16, pady=(6, 4))

        self._search_btn = tk.Button(
            row2, text="  SEARCH  ",
            bg=ACCENT, fg=WHITE,
            font=("Courier New", 10, "bold"), relief='flat',
            padx=14, pady=6, cursor='hand2',
            command=self._do_search)
        self._search_btn.pack(side='left')
        self._search_btn.bind("<Return>", lambda e: self._do_search())

        tk.Button(row2, text="CLEAR RESULTS",
                  bg=PANEL_BG, fg=TEXT_DIM,
                  font=("Courier New", 9), relief='flat',
                  padx=10, pady=6, cursor='hand2',
                  command=self._clear_results).pack(side='left', padx=8)

        tk.Button(row2, text="REBUILD INDEX",
                  bg=PANEL_BG, fg=TEXT_DIM,
                  font=("Courier New", 9), relief='flat',
                  padx=10, pady=6, cursor='hand2',
                  command=self._rebuild_index).pack(side='left', padx=0)

        tk.Button(row2, text="  DOWNLOAD RESULTS  ",
                  bg=GREEN, fg=WHITE,
                  font=("Courier New", 9, "bold"), relief='flat',
                  padx=10, pady=6, cursor='hand2',
                  command=self._download_results).pack(side='right', padx=(6, 0))

        self._status_lbl = tk.Label(row2, text="",
                                     bg=PANEL_BG, fg=TEXT_DIM,
                                     font=("Courier New", 9))
        self._status_lbl.pack(side='left', padx=8)

        # FIX 6: Updated column headers to show "YEAR / SOURCE" instead of just "SOURCE"
        col_hdr = tk.Frame(self, bg=CARD_BG)
        col_hdr.pack(fill='x', pady=(0, 2))
        for col, w in [("#", 3), ("TYPE", 8), ("SERIAL", 10), ("LAB NUMBER", 16),
                        ("CERT DATE", 12), ("SHIPPED", 12), ("CHPS", 8), ("ADDR", 8), ("SOURCE FILE", 20)]:
            tk.Label(col_hdr, text=col, bg=CARD_BG, fg=ACCENT,
                     font=("Courier New", 8, "bold"),
                     width=w, anchor='w').pack(side='left', padx=4, pady=4)

        results_outer = tk.Frame(self, bg=BG)
        results_outer.pack(fill='both', expand=True)

        self._canvas = tk.Canvas(results_outer, bg=BG, highlightthickness=0)
        sb = ttk.Scrollbar(results_outer, orient='vertical',
                           command=self._canvas.yview)
        self._canvas.configure(yscrollcommand=sb.set)
        sb.pack(side='right', fill='y')
        self._canvas.pack(side='left', fill='both', expand=True)

        self._results_frame = tk.Frame(self._canvas, bg=BG)
        self._win = self._canvas.create_window(
            (0, 0), window=self._results_frame, anchor='nw')
        self._results_frame.bind('<Configure>',
            lambda e: self._canvas.configure(
                scrollregion=self._canvas.bbox('all')))
        self._canvas.bind('<Configure>',
            lambda e: self._canvas.itemconfig(self._win, width=e.width))

        self._result_count = 0
        self._results_data = []

        self._index_bar = tk.Label(self, text="Search index: building...",
                                    bg=PANEL_BG, fg=AMBER,
                                    font=("Courier New", 8), anchor='w')
        self._index_bar.pack(fill='x', padx=8, pady=(2, 0))

    def set_index_status(self, msg):
        try:
            if 'ready' in msg.lower():
                self._index_bar.config(text=f"✓ {msg}", fg=GREEN)
            elif 'building' in msg.lower() or 'indexing' in msg.lower():
                self._index_bar.config(text=f"⟳ {msg}", fg=AMBER)
            else:
                self._index_bar.config(text=msg, fg=TEXT_DIM)
        except Exception:
            pass

    def focus_first(self):
        self._num_entry.focus_set()

    def _clear_results(self):
        for w in self._results_frame.winfo_children():
            w.destroy()
        self._result_count = 0
        self._results_data = []
        self._status_lbl.config(text="Results cleared.")

    def _rebuild_index(self):
        self.app._start_index_build()
        self._status_lbl.config(text="Rebuilding index...", fg=AMBER)

    def _do_search(self):
        try:
            n = int(self._num_var.get().strip())
            if n < 1:
                raise ValueError
        except ValueError:
            messagebox.showwarning("Invalid Input",
                                   "Enter a valid number of units (1 or more).")
            return

        type_label = self._type_var.get()
        serial     = self._serial_var.get().strip()

        if not serial:
            messagebox.showwarning("Missing Serial",
                                   "Enter a serial number to search.")
            self._serial_entry.focus_set()
            return

        cat, utype = self._type_map[type_label]
        sheet_name = cat

        self._status_lbl.config(text="Searching all years...", fg=TEXT_DIM)
        self.update_idletasks()

        # FIX 6: Search ALL entries across ALL years for this serial
        all_records = search_all_entries_for_serial(serial, sheet_name)

        if all_records:
            # Add a section header for this serial
            self._add_serial_header(utype, serial, len(all_records))
            for rec in all_records:
                self._add_result_row(utype, serial, rec)
            self._status_lbl.config(
                text=f"Found {len(all_records)} record(s). Total rows: {self._result_count}", fg=GREEN)
        else:
            self._add_not_found_row(utype, serial)
            self._status_lbl.config(
                text=f"Not found in any year. Total rows: {self._result_count}", fg=AMBER)

        if n > 1:
            self._multi_search(n - 1, cat, sheet_name)

        self._canvas.update_idletasks()
        self._canvas.yview_moveto(1.0)

    def _add_serial_header(self, utype, serial, count):
        """FIX 6: Visual separator/header when showing multiple records for one serial."""
        hdr = tk.Frame(self._results_frame, bg=CARD_BG)
        hdr.pack(fill='x', pady=(6, 0))
        tk.Label(hdr,
                 text=f"  ▸ {utype} serial {serial} — {count} record(s) across all years",
                 bg=CARD_BG, fg=ACCENT,
                 font=("Courier New", 9, "bold"), anchor='w').pack(side='left', padx=6, pady=3)

    def _search_unit(self, serial, sheet_name, utype):
        """Returns most recent single record (used for multi-unit searches)."""
        return search_all_history(serial, sheet_name)

    def _add_result_row(self, utype, serial, rec):
        self._result_count += 1
        i    = self._result_count
        bg   = BG if i % 2 == 0 else PANEL_BG
        self._results_data.append({
            'index':       i,
            'type':        utype,
            'serial':      serial,
            'lab_number':  str(rec.get('lab_number', '—')),
            'cert_date':   str(rec.get('date', '—')).split(" ")[0] if rec.get('date') else '—',
            'shipped_date':str(rec.get('shipped_date', '')).split(" ")[0] if rec.get('shipped_date') else '—',
            'chps':        str(rec.get('chps_number', '—')),
            'address':     str(rec.get('address_code', '—')),
            'source':      str(rec.get('source_file', '—')),
            'found':       True,
        })

        date_str = str(rec['date']).split(" ")[0] if rec['date'] else "—"
        yrs      = years_since(rec['date'])
        yrs_str  = f"{yrs:.1f}y ago" if yrs is not None else ""
        flag_fg  = RED if (yrs is not None and yrs < 3.0) else GREEN

        row = tk.Frame(self._results_frame, bg=bg)
        row.pack(fill='x', pady=1)

        ship_raw  = rec.get('shipped_date', None)
        ship_str  = str(ship_raw).split(" ")[0] if ship_raw else "—"
        ship_fg   = GREEN if ship_raw else TEXT_DIM

        for val, w, color in [
            (str(i),                     3,  TEXT),
            (utype,                      8,  TEXT),
            (serial,                    10,  TEXT),
            (str(rec['lab_number']),   16,  TEXT),
            (date_str,                  12,  TEXT),
            (ship_str,                  12,  ship_fg),
            (str(rec['chps_number']),  8,  TEXT),
            (str(rec['address_code']), 8,  TEXT),
            (str(rec['source_file']), 20,  TEXT_DIM),
        ]:
            tk.Label(row, text=val, bg=bg, fg=color,
                     font=("Courier New", 9),
                     width=w, anchor='w').pack(side='left', padx=4, pady=4)

        tk.Label(row, text=yrs_str, bg=bg, fg=flag_fg,
                 font=("Courier New", 8, "bold")).pack(side='left', padx=4)

        if yrs is not None and yrs < 3.0:
            tk.Label(row, text=f"⚠ UNDER 3 YRS", bg=bg, fg=RED,
                     font=("Courier New", 8, "bold")).pack(side='left', padx=4)

    def _add_not_found_row(self, utype, serial):
        self._result_count += 1
        i  = self._result_count
        bg = BG if i % 2 == 0 else PANEL_BG
        self._results_data.append({
            'index': i, 'type': utype, 'serial': serial,
            'lab_number': '—', 'cert_date': '—', 'shipped_date': '—',
            'chps': '—', 'address': '—', 'source': 'NOT FOUND', 'found': False,
        })
        row = tk.Frame(self._results_frame, bg=bg)
        row.pack(fill='x', pady=1)
        for val, w in [(str(i), 3), (utype, 8), (serial, 10)]:
            tk.Label(row, text=val, bg=bg, fg=TEXT,
                     font=("Courier New", 9),
                     width=w, anchor='w').pack(side='left', padx=4, pady=4)
        tk.Label(row, text="NOT FOUND IN ANY HISTORY FILE",
                 bg=bg, fg=TEXT_DIM,
                 font=("Courier New", 9, "italic")).pack(side='left', padx=4)

    def _download_results(self):
        if not self._results_data:
            messagebox.showwarning("No Results", "No search results to download.")
            return

        today     = datetime.now()
        date_str  = today.strftime("%m/%d/%Y")
        fname     = f"STATUS_REQUEST_{today.strftime('%m%d%Y')}.docx"
        out_path  = os.path.join(BASE_DIR, fname)

        try:
            doc = docx.Document()

            title = doc.add_heading("RADAR / LIDAR UNIT STATUS REPORT", level=1)
            title.alignment = 1
            doc.add_paragraph(f"Date of Request: {date_str}").alignment = 1
            doc.add_paragraph("")

            headers = ["#", "TYPE", "SERIAL", "LAB NUMBER",
                       "CERT DATE", "SHIPPED DATE", "CHPS", "ADDRESS CODE", "SOURCE FILE"]
            tbl = doc.add_table(rows=1, cols=len(headers))
            tbl.style = 'Table Grid'

            hdr_row = tbl.rows[0]
            for i, h in enumerate(headers):
                cell = hdr_row.cells[i]
                cell.text = h
                run = cell.paragraphs[0].runs[0]
                run.bold = True
                from docx.shared import RGBColor, Pt
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                tc  = cell._tc
                tcp = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'),   'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'),  'D41736')
                tcp.append(shd)

            from docx.shared import Pt, RGBColor
            for rec in self._results_data:
                row = tbl.add_row()
                vals = [
                    str(rec['index']),
                    rec['type'],
                    rec['serial'],
                    rec['lab_number'],
                    rec['cert_date'],
                    rec['shipped_date'],
                    rec['chps'],
                    rec['address'],
                    rec['source'],
                ]
                for i, v in enumerate(vals):
                    cell = row.cells[i]
                    cell.text = v
                    run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(v)
                    run.font.size = Pt(9)
                    if not rec['found']:
                        run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
                    elif i == 5 and rec['shipped_date'] != '—':
                        run.font.color.rgb = RGBColor(0x00, 0xA3, 0x9D)

            doc.add_paragraph("")
            found_count = sum(1 for r in self._results_data if r['found'])
            not_found   = len(self._results_data) - found_count
            shipped     = sum(1 for r in self._results_data if r['found'] and r['shipped_date'] != '—')
            summary = doc.add_paragraph()
            summary.add_run(
                f"Total units searched: {len(self._results_data)}    "
                f"Found: {found_count}    "
                f"Not found: {not_found}    "
                f"Shipped: {shipped}"
            ).font.size = Pt(9)

            doc.save(out_path)
            messagebox.showinfo(
                "Downloaded",
                f"Status report saved as:\n  {fname}\n\n"
                f"Units: {len(self._results_data)}  |  Found: {found_count}  |  Shipped: {shipped}")

        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate report:\n{e}")

    def _multi_search(self, remaining, cat, sheet_name):
        """For n>1 searches, open dialog to enter each additional serial.
        FIX 6: Each serial shows ALL records across all years."""
        for _ in range(remaining):
            dlg = tk.Toplevel(self.winfo_toplevel())
            dlg.title(f"Search Unit")
            dlg.configure(bg=WHITE)
            dlg.resizable(False, False)
            dlg.grab_set()
            dlg.update_idletasks()
            sw = dlg.winfo_screenwidth(); sh = dlg.winfo_screenheight()
            dlg.geometry(f"380x200+{(sw-380)//2}+{(sh-200)//2}")

            tk.Label(dlg, text=f"Next Unit — {cat}",
                     bg=WHITE, fg=ACCENT,
                     font=("Courier New", 11, "bold")).pack(pady=(16, 4))

            tk.Label(dlg, text="Serial Number:",
                     bg=WHITE, fg=TEXT,
                     font=("Courier New", 10)).pack(anchor='w', padx=20)
            sv = tk.StringVar()
            se = tk.Entry(dlg, textvariable=sv,
                          bg=PANEL_BG, fg=TEXT, insertbackground=TEXT,
                          font=("Courier New", 13, "bold"),
                          width=16, relief='solid', bd=1)
            se.pack(padx=20, pady=6, ipady=5)
            se.focus_set()

            result = [None]
            def on_ok(sv=sv, r=result, d=dlg):
                r[0] = sv.get().strip(); d.destroy()
            def on_skip(r=result, d=dlg):
                r[0] = '__skip__'; d.destroy()

            bf = tk.Frame(dlg, bg=WHITE)
            bf.pack(pady=8)
            ob = tk.Button(bf, text="SEARCH", bg=ACCENT, fg=WHITE,
                      font=("Courier New", 10, "bold"), relief='flat',
                      padx=14, pady=6, cursor='hand2', command=on_ok)
            ob.pack(side='left', padx=6)
            ob.bind("<Return>", lambda e, f=on_ok: f())
            se.bind("<Return>", lambda e, f=on_ok: f())
            tk.Button(bf, text="SKIP", bg=PANEL_BG, fg=TEXT_DIM,
                      font=("Courier New", 10), relief='flat',
                      padx=14, pady=6, cursor='hand2',
                      command=on_skip).pack(side='left', padx=6)
            dlg.protocol("WM_DELETE_WINDOW", on_skip)
            self.winfo_toplevel().wait_window(dlg)

            if result[0] and result[0] != '__skip__':
                serial = result[0]
                utype_display = self._type_var.get().split(": ")[-1]
                # FIX 6: Get ALL records for this serial too
                all_recs = search_all_entries_for_serial(serial, sheet_name)
                if all_recs:
                    self._add_serial_header(utype_display, serial, len(all_recs))
                    for rec in all_recs:
                        self._add_result_row(utype_display, serial, rec)
                else:
                    self._add_not_found_row(utype_display, serial)
                self._canvas.update_idletasks()
                self._canvas.yview_moveto(1.0)


# ─────────────────────────────────────────────
#  RUN
# ─────────────────────────────────────────────
if __name__ == "__main__":
    app = EntryApp()
    app.mainloop()